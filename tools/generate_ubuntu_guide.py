#!/usr/bin/env python3
"""Generate QEMU-VCS CoSim Platform Ubuntu Usage Guide (.docx)."""

import os
import sys

sys.path.insert(0, os.path.expanduser("~/.local/lib/python3.12/site-packages"))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Output path
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_PATH = os.path.join(PROJECT_DIR, "docs", "CoSim-Platform-Ubuntu-Usage-Guide.docx")


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def make_table(doc, headers, rows):
    """Create a formatted table with header styling."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)

    return table


def add_code_block(doc, code_text):
    """Add a code block with monospace font and gray shading."""
    for line in code_text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        # Set shading on paragraph
        pPr = p._p.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F2F2F2"/>')
        pPr.append(shd)


def add_para(doc, text, bold=False, size=10):
    """Add a simple paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def generate():
    doc = Document()

    # Set A4 page size
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ========== TITLE PAGE ==========
    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("QEMU-VCS \u534f\u540c\u4eff\u771f\u5e73\u53f0\n\u4f7f\u7528\u8bf4\u660e\u4e66\uff08Ubuntu \u7248\uff09")
    run.bold = True
    run.font.size = Pt(24)

    doc.add_paragraph()

    info_lines = [
        "\u7248\u672c\uff1aP5\uff08VIP \u6a21\u5f0f + ETH \u6027\u80fd\u57fa\u51c6\uff09",
        "\u65e5\u671f\uff1a2026-04-21",
        "\u7ef4\u62a4\u8005\uff1aBeihang-yuting <2965455908@qq.com>",
        "\u4ed3\u5e93\uff1ahttps://github.com/Beihang-yuting/QEMU_VCS",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)

    doc.add_page_break()

    # ========== 1. \u5e73\u53f0\u6982\u8ff0 ==========
    doc.add_heading("1. \u5e73\u53f0\u6982\u8ff0", level=1)

    doc.add_heading("1.1 \u6838\u5fc3\u80fd\u529b", level=2)
    capabilities = [
        "PCIe MMIO \u8bfb\u5199\u4e8b\u52a1\u8054\u5408\u4eff\u771f\uff08TLP \u7ea7\uff09",
        "DMA \u6570\u636e\u9762\u53cc\u5411\u4f20\u8f93\uff08\u8bbe\u5907\u2192Host / Host\u2192\u8bbe\u5907\uff09",
        "MSI \u4e2d\u65ad\u6ce8\u5165\uff08VCS RTL \u2192 QEMU Guest\uff09",
        "\u53cc\u6a21\u5f0f\u540c\u6b65\uff1a\u5feb\u901f\u6a21\u5f0f\uff08\u4e8b\u52a1\u7ea7\uff09\u4e0e\u7cbe\u786e\u6a21\u5f0f\uff08\u5468\u671f\u7ea7\u9501\u6b65\uff09",
        "\u8fd0\u884c\u65f6\u6a21\u5f0f\u5207\u6362\u3001\u4e8b\u52a1\u7ea7\u8ffd\u8e2a\u65e5\u5fd7\uff08CSV/JSON\uff09",
        "POSIX \u5171\u4eab\u5185\u5b58 + Unix Domain Socket \u4f4e\u5f00\u9500 IPC",
        "\u53cc\u8282\u70b9\u4ee5\u592a\u7f51\u4e92\u6253\uff08P3\uff09\uff1aETH SHM \u5e27\u961f\u5217\u30019KB Jumbo\u3001\u677e\u8026\u5408\u65f6\u95f4\u540c\u6b65",
        "\u94fe\u8def\u6a21\u578b\uff08P3\uff09\uff1a\u4e22\u5305\u7387\u3001\u7a81\u53d1\u4e22\u5305\u3001\u56fa\u5b9a\u5ef6\u8fdf\u3001\u901f\u7387\u9650\u3001\u6d41\u63a7\u7a97\u53e3",
        "\u8c03\u8bd5\u5de5\u5177\u94fe\uff08P4\uff09\uff1acosim_cli REPL\u3001trace_analyzer\u3001launch_dual \u7f16\u6392",
        "UVM VIP \u6a21\u5f0f\uff08P5\uff09\uff1aPCIe TL VIP \u9a8c\u8bc1\u73af\u5883\uff0c\u652f\u6301 UVM factory override\u3001cosim_rc_driver\u3001completion \u901a\u8fc7 VIP \u56de\u4f20",
        "ETH \u6027\u80fd\u57fa\u51c6\uff08P5\uff09\uff1aiperf \u98ce\u683c\u541e\u5410\u91cf\u6d4b\u8bd5\u3001\u53cc\u5411\u4e92\u6253\u6d4b\u8bd5\u3001Ring Depth \u8c03\u4f18",
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style="List Bullet")

    doc.add_heading("1.2 \u7cfb\u7edf\u67b6\u6784", level=2)
    arch_items = [
        "QEMU \u8fdb\u7a0b\uff1a\u8fd0\u884c Guest Linux\uff0c\u52a0\u8f7d\u81ea\u5b9a\u4e49 PCIe RC \u8bbe\u5907 cosim-pcie-rc",
        "VCS \u4eff\u771f\u8fdb\u7a0b\uff1a\u8fd0\u884c RTL \u8bbe\u8ba1\uff08DPU PCIe EP\uff09\uff0c\u901a\u8fc7 DPI-C \u8c03\u7528 Bridge",
        "Bridge \u5e93\uff1a\u627f\u8f7d TLP \u8bf7\u6c42\u961f\u5217\u3001\u5b8c\u6210\u961f\u5217\u3001DMA \u6570\u636e\u533a\u3001MSI \u4e8b\u4ef6\u961f\u5217\u3001\u65f6\u949f\u540c\u6b65\u63e1\u624b",
    ]
    for item in arch_items:
        doc.add_paragraph(item, style="List Bullet")

    # ========== 2. \u4ed3\u5e93\u76ee\u5f55\u7ed3\u6784 ==========
    doc.add_heading("2. \u4ed3\u5e93\u76ee\u5f55\u7ed3\u6784", level=1)
    dir_tree = """\
cosim-platform/
\u251c\u2500\u2500 CMakeLists.txt              # \u9876\u5c42 CMake
\u251c\u2500\u2500 Makefile                    # \u4e00\u952e make / make test + VCS \u7f16\u8bd1\u76ee\u6807
\u251c\u2500\u2500 config.env                  # \u73af\u5883\u914d\u7f6e\u6587\u4ef6
\u251c\u2500\u2500 setup.sh                    # \u4e00\u952e\u5b89\u88c5\u811a\u672c
\u251c\u2500\u2500 cosim.sh                    # \u7edf\u4e00 CLI \u5de5\u5177
\u251c\u2500\u2500 bridge/                     # Bridge \u5e93\u6e90\u7801
\u2502   \u251c\u2500\u2500 common/                 # \u5e73\u53f0\u5171\u7528\uff1aSHM\u3001\u73af\u5f62\u7f13\u51b2\u3001DMA\u3001trace\u3001ETH SHM\u3001\u94fe\u8def\u6a21\u578b
\u2502   \u251c\u2500\u2500 qemu/                   # QEMU \u4fa7 Bridge\uff08libcosim_bridge.so\uff09
\u2502   \u251c\u2500\u2500 vcs/                    # VCS \u4fa7 Bridge\uff08DPI-C + .sv\uff09
\u2502   \u2514\u2500\u2500 eth/                    # \u4ee5\u592a\u7f51\u7aef\u53e3 + MAC stub + MAC DPI
\u251c\u2500\u2500 qemu-plugin/                # QEMU \u81ea\u5b9a\u4e49 PCIe RC \u8bbe\u5907
\u251c\u2500\u2500 pcie_tl_vip/src/            # PCIe TL VIP\uff08P5 \u65b0\u589e\uff09
\u2502   \u251c\u2500\u2500 pcie_tl_if.sv           # 256-bit \u5355\u5411 TLP \u603b\u7ebf\u63a5\u53e3
\u2502   \u2514\u2500\u2500 pcie_tl_pkg.sv          # VIP \u7c7b\u578b/\u7c7b\u5b9a\u4e49
\u251c\u2500\u2500 vcs-tb/                     # VCS \u6d4b\u8bd5\u5e73\u53f0
\u2502   \u251c\u2500\u2500 tb_top.sv               # Legacy \u6a21\u5f0f\u9876\u5c42
\u2502   \u251c\u2500\u2500 pcie_ep_stub.sv         # PCIe EP \u7b80\u5316\u6a21\u578b
\u2502   \u251c\u2500\u2500 cosim_vip_top.sv        # VIP \u6a21\u5f0f\u9876\u5c42\uff08P5\uff09
\u2502   \u251c\u2500\u2500 cosim_pkg.sv            # UVM \u7ec4\u4ef6\u5305\uff08P5\uff09
\u2502   \u251c\u2500\u2500 cosim_test.sv           # UVM test\uff08P5\uff09
\u2502   \u251c\u2500\u2500 cosim_rc_driver.sv      # cosim RC driver - factory override\uff08P5\uff09
\u2502   \u2514\u2500\u2500 glue_if_to_stub.sv      # VIP \u2194 Stub \u4fe1\u53f7\u8f6c\u6362\uff08P5\uff09
\u251c\u2500\u2500 uvm-tb/                     # \u72ec\u7acb UVM TB\uff08build_uvm.sh \u7f16\u8bd1\uff09
\u251c\u2500\u2500 scripts/                    # \u6d4b\u8bd5/\u542f\u52a8/\u5de5\u5177\u811a\u672c
\u251c\u2500\u2500 tests/
\u2502   \u251c\u2500\u2500 unit/                   # \u5355\u5143\u6d4b\u8bd5
\u2502   \u251c\u2500\u2500 integration/            # \u96c6\u6210\u6d4b\u8bd5
\u2502   \u2514\u2500\u2500 e2e/                    # \u7aef\u5230\u7aef\u6d4b\u8bd5\uff08P5 \u65b0\u589e\uff09
\u2502       \u251c\u2500\u2500 test_vip_smoke.c    # VIP \u5192\u70df\u6d4b\u8bd5
\u2502       \u251c\u2500\u2500 test_eth_iperf.c    # ETH \u5355\u5411\u541e\u5410\u91cf\u6d4b\u8bd5
\u2502       \u2514\u2500\u2500 test_eth_bidir.c    # ETH \u53cc\u5411\u4e92\u6253\u6d4b\u8bd5
\u251c\u2500\u2500 tools/                      # eth_tap_bridge\u3001\u6587\u6863\u751f\u6210\u5de5\u5177
\u251c\u2500\u2500 images/                     # \u6784\u5efa\u4ea7\u7269\uff08\u5185\u6838\u3001initramfs\uff09
\u2514\u2500\u2500 docs/                       # \u4f7f\u7528\u8bf4\u660e\u4e0e\u8c03\u8bd5\u6307\u5357"""
    add_code_block(doc, dir_tree)

    # ========== 3. \u73af\u5883\u8981\u6c42 ==========
    doc.add_heading("3. \u73af\u5883\u8981\u6c42", level=1)

    doc.add_heading("3.1 \u64cd\u4f5c\u7cfb\u7edf\u4e0e\u57fa\u7840\u5de5\u5177", level=2)
    make_table(doc,
        ["\u7ec4\u4ef6", "\u7248\u672c\u8981\u6c42", "\u7528\u9014"],
        [
            ["Linux", "Ubuntu 20.04 / 22.04 / 24.04 (x86_64)", "\u8fd0\u884c QEMU + VCS"],
            ["GCC", ">= 9.0\uff08\u652f\u6301 C11 / __atomic\uff09", "\u7f16\u8bd1 Bridge"],
            ["CMake", ">= 3.16", "\u6784\u5efa\u7cfb\u7edf"],
            ["Python3", ">= 3.8", "QEMU \u6784\u5efa + \u5de5\u5177\u811a\u672c"],
            ["Meson + Ninja", "-", "QEMU \u6784\u5efa\u4f9d\u8d56"],
            ["pthread / librt", "glibc \u5185\u7f6e", "Bridge \u7ebf\u7a0b\u4e0e SHM"],
            ["QEMU \u6e90\u7801", "9.2.0", "\u8fd0\u884c Guest OS"],
            ["Synopsys VCS", "Q-2020.03+\uff08\u542b DPI-C\uff09", "RTL \u4eff\u771f"],
        ])

    doc.add_heading("3.2 Ubuntu \u7cfb\u7edf\u4f9d\u8d56\u5b89\u88c5", level=2)
    add_code_block(doc, """\
sudo apt update
sudo apt install -y \\
    gcc g++ make cmake git \\
    meson ninja-build pkg-config \\
    libglib2.0-dev libpixman-1-dev libslirp-dev \\
    python3 python3-pip python3-venv \\
    cpio gzip""")

    doc.add_heading("3.3 \u53ef\u9009\u7ec4\u4ef6", level=2)
    optional = [
        "KVM\uff08/dev/kvm\uff09\uff1a\u542f\u7528\u540e Guest \u8fd0\u884c\u6027\u80fd\u63a5\u8fd1\u539f\u751f",
        "GDB\uff1a\u8c03\u8bd5 Guest \u5185\u6838\u6216\u9a71\u52a8",
        "VCS \u8bb8\u53ef\u8bc1\uff1a\u786e\u4fdd SNPSLMD_LICENSE_FILE \u6307\u5411\u6b63\u786e\u7684 license \u6587\u4ef6",
    ]
    for item in optional:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.4 \u8d44\u6e90\u9700\u6c42", level=2)
    resources = [
        "\u78c1\u76d8\uff1a\u7ea6 500 MB\uff08Bridge + QEMU \u6784\u5efa\u4ea7\u7269\uff0c\u4e0d\u542b Guest \u955c\u50cf\uff09",
        "\u5185\u5b58\uff1a\u81f3\u5c11 8 GB\uff08\u542b QEMU Guest 4 GB + VCS 2 GB\uff09",
        "\u5171\u4eab\u5185\u5b58\uff1a/dev/shm \u81f3\u5c11 64 MB\uff08Ubuntu \u9ed8\u8ba4\u6302\u8f7d\u4e3a tmpfs\uff0c\u5927\u5c0f\u4e3a\u7269\u7406\u5185\u5b58\u4e00\u534a\uff09",
    ]
    for item in resources:
        doc.add_paragraph(item, style="List Bullet")

    # ========== 4. \u5feb\u901f\u5f00\u59cb ==========
    doc.add_heading("4. \u5feb\u901f\u5f00\u59cb", level=1)

    doc.add_heading("4.1 \u514b\u9686\u4ed3\u5e93", level=2)
    add_code_block(doc, """\
git clone https://github.com/Beihang-yuting/QEMU_VCS.git cosim-platform
cd cosim-platform""")

    doc.add_heading("4.2 \u914d\u7f6e\u73af\u5883\u53d8\u91cf", level=2)
    add_para(doc, "\u7f16\u8f91 config.env \u6587\u4ef6\uff1a")
    make_table(doc,
        ["\u53d8\u91cf\u540d", "\u9ed8\u8ba4\u503c", "\u8bf4\u660e"],
        [
            ["VCS_HOME", "(\u7a7a)", "VCS \u5b89\u88c5\u76ee\u5f55\uff0c\u7559\u7a7a\u81ea\u52a8\u641c\u7d22"],
            ["SNPSLMD_LICENSE_FILE", "/opt/synopsys/license/license.dat", "VCS \u8bb8\u53ef\u8bc1\u8def\u5f84"],
            ["QEMU_VERSION", "v9.2.0", "QEMU \u7248\u672c\u53f7"],
            ["GUEST1_IP / GUEST2_IP", "10.0.0.1 / 10.0.0.2", "Guest \u7f51\u7edc IP"],
            ["GUEST_MEMORY", "256M", "Guest \u5185\u5b58\u5927\u5c0f"],
        ])

    doc.add_heading("4.3 \u4e00\u952e\u5b89\u88c5", level=2)
    add_code_block(doc, "bash setup.sh")
    add_para(doc, "setup.sh \u5b89\u88c5\u6b65\u9aa4\u8bf4\u660e\uff1a")
    steps = [
        "[1/9] \u52a0\u8f7d\u914d\u7f6e\u6587\u4ef6",
        "[2/9] \u68c0\u6d4b Ubuntu \u7cfb\u7edf\u5e76\u5b89\u88c5\u4f9d\u8d56\uff08apt install\uff09",
        "[3/9] \u7f16\u8bd1 Bridge \u5e93\uff08libcosim_bridge.so\uff09",
        "[4/9] \u7f16\u8bd1 QEMU\uff08\u6ce8\u5165 cosim_pcie_rc \u8bbe\u5907\uff09",
        "[5/9] \u7f16\u8bd1 VCS simv\uff08\u542b DPI-C bridge\uff09",
        "[6/9] \u7f16\u8bd1 eth_tap_bridge",
        "[7/9] \u6784\u5efa initramfs",
        "[8/9] \u8fd0\u884c\u5355\u5143\u6d4b\u8bd5",
        "[9/9] \u5b89\u88c5\u6458\u8981",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Bullet")

    add_para(doc, "\u6784\u5efa\u4ea7\u7269\uff1a", bold=True)
    make_table(doc,
        ["\u4ea7\u7269", "\u8def\u5f84", "\u8bf4\u660e"],
        [
            ["libcosim_bridge.so", "build/bridge/", "SHM \u901a\u4fe1\u6838\u5fc3\u5e93"],
            ["qemu-system-x86_64", "third_party/qemu/build/", "\u542b cosim-pcie-rc \u8bbe\u5907\u7684 QEMU"],
            ["simv", "vcs-tb/sim_build/", "VCS \u4eff\u771f\u53ef\u6267\u884c\u6587\u4ef6"],
            ["eth_tap_bridge", "tools/", "TAP \u6865\u63a5\u5de5\u5177"],
            ["initramfs-*.cpio.gz", "images/", "Guest \u542f\u52a8\u955c\u50cf"],
        ])

    doc.add_heading("4.4 \u624b\u52a8\u6784\u5efa\uff08\u4e0d\u4f7f\u7528 setup.sh\uff09", level=2)
    add_code_block(doc, """\
# \u6784\u5efa Bridge
make bridge

# \u8fd0\u884c\u5168\u90e8\u6d4b\u8bd5
make test

# \u96c6\u6210 QEMU
export QEMU_SRC=./third_party/qemu
./scripts/setup_cosim_qemu.sh $QEMU_SRC
cd $QEMU_SRC
./configure --target-list=x86_64-softmmu --enable-kvm
make -j$(nproc)""")

    # ========== 5. VCS \u7f16\u8bd1 ==========
    doc.add_heading("5. VCS \u7f16\u8bd1\uff08Legacy / VIP \u6a21\u5f0f\uff09", level=1)

    doc.add_heading("5.1 Legacy \u6a21\u5f0f", level=2)
    add_code_block(doc, """\
make vcs-legacy
# \u6216\u624b\u52a8:
vcs -full64 -sverilog -timescale=1ns/1ps \\
    -CFLAGS "-I bridge/common -I bridge/vcs -I bridge/qemu -std=c99" \\
    -LDFLAGS "-lrt -lpthread" \\
    bridge/vcs/bridge_vcs.sv vcs-tb/tb_top.sv vcs-tb/pcie_ep_stub.sv \\
    bridge/vcs/bridge_vcs.c bridge/common/shm_layout.c \\
    bridge/common/ring_buffer.c bridge/common/dma_manager.c \\
    bridge/common/trace_log.c bridge/common/eth_shm.c \\
    bridge/common/link_model.c bridge/vcs/vq_eth_stub.c \\
    bridge/vcs/sock_sync_vcs.c \\
    -o build/simv_legacy""")

    doc.add_heading("5.2 VIP \u6a21\u5f0f\uff08P5 \u65b0\u589e\uff09", level=2)
    add_code_block(doc, """\
make vcs-vip
# \u6216\u624b\u52a8:
vcs -full64 -sverilog -timescale=1ns/1ps -ntb_opts uvm-1.2 \\
    +define+COSIM_VIP_MODE \\
    -CFLAGS "-I $(pwd)/bridge/common -I $(pwd)/bridge/vcs -I $(pwd)/bridge/qemu -std=c99" \\
    -LDFLAGS "-lrt -lpthread" \\
    +incdir+bridge/vcs +incdir+pcie_tl_vip/src +incdir+vcs-tb \\
    bridge/vcs/bridge_vcs.sv \\
    pcie_tl_vip/src/pcie_tl_if.sv pcie_tl_vip/src/pcie_tl_pkg.sv \\
    vcs-tb/tb_top.sv vcs-tb/pcie_ep_stub.sv \\
    vcs-tb/glue_if_to_stub.sv vcs-tb/cosim_pkg.sv \\
    vcs-tb/cosim_vip_top.sv \\
    bridge/vcs/bridge_vcs.c bridge/common/shm_layout.c \\
    bridge/common/ring_buffer.c bridge/common/dma_manager.c \\
    bridge/common/trace_log.c bridge/common/eth_shm.c \\
    bridge/common/link_model.c bridge/vcs/vq_eth_stub.c \\
    bridge/vcs/sock_sync_vcs.c \\
    -o build/simv_vip""")

    doc.add_heading("5.3 VIP + \u6027\u80fd\u7edf\u8ba1\u6a21\u5f0f", level=2)
    add_code_block(doc, "make vcs-vip-perf")

    # ========== 6. \u542f\u52a8\u8054\u5408\u4eff\u771f ==========
    doc.add_heading("6. \u542f\u52a8\u8054\u5408\u4eff\u771f", level=1)

    doc.add_heading("6.1 \u542f\u52a8 QEMU\uff08\u7ec8\u7aef A\uff09", level=2)
    add_code_block(doc, """\
export GUEST_KERNEL=/path/to/bzImage
export GUEST_ROOTFS=/path/to/rootfs.qcow2   # \u53ef\u9009
./scripts/run_cosim.sh""")

    doc.add_heading("6.2 \u542f\u52a8 VCS\uff08\u7ec8\u7aef B\uff09", level=2)
    add_para(doc, "Legacy \u6a21\u5f0f\uff1a", bold=True)
    add_code_block(doc, "./build/simv_legacy +SHM_NAME=/cosim0 +SOCK_PATH=/tmp/cosim.sock")
    add_para(doc, "VIP \u6a21\u5f0f\uff1a", bold=True)
    add_code_block(doc, "./build/simv_vip +SHM_NAME=/cosim0 +SOCK_PATH=/tmp/cosim.sock +UVM_TESTNAME=cosim_test")
    add_para(doc, "\u6ce8\u610f\uff1aQEMU \u5fc5\u987b\u5148\u542f\u52a8\uff08server \u7aef\uff09\uff0cVCS \u4f5c\u4e3a client \u63a5\u5165\u3002\u9ed8\u8ba4 SHM_NAME=/cosim0, SOCK_PATH=/tmp/cosim.sock\u3002")

    # ========== 7. \u6d4b\u8bd5\u6d41\u7a0b ==========
    doc.add_heading("7. \u6d4b\u8bd5\u6d41\u7a0b", level=1)

    doc.add_heading("7.1 \u6d4b\u8bd5\u9636\u6bb5\u603b\u89c8", level=2)
    make_table(doc,
        ["\u9636\u6bb5", "\u63cf\u8ff0", "\u9a8c\u8bc1\u5185\u5bb9", "\u547d\u4ee4"],
        [
            ["Phase 1", "Config Space", "PCIe \u914d\u7f6e\u7a7a\u95f4\u8bfb\u5199", "./cosim.sh test phase1"],
            ["Phase 2", "DMA \u8bfb\u5199", "Host-Guest DMA \u6570\u636e\u4f20\u8f93", "./cosim.sh test phase2"],
            ["Phase 3", "MSI \u4e2d\u65ad", "MSI \u4e2d\u65ad\u89e6\u53d1\u4e0e\u63a5\u6536", "./cosim.sh test phase3"],
            ["Phase 4", "\u73af\u56de\u6d4b\u8bd5", "VirtIO-Net + DMA + \u53cc VCS \u73af\u56de", "./cosim.sh test phase4"],
            ["Phase 5", "\u7aef\u5230\u7aef\u7f51\u7edc", "Guest \u95f4 ping/nc/iperf3", "./cosim.sh test phase5"],
            ["TAP", "\u5bbf\u4e3b\u673a\u6865\u63a5", "Guest \u4e0e Host \u7f51\u7edc\u901a\u4fe1", "./cosim.sh test tap"],
        ])

    doc.add_heading("7.2 \u5355\u5143 + \u96c6\u6210\u6d4b\u8bd5", level=2)
    add_code_block(doc, "make test")
    add_para(doc, "\u5171 17 \u9879\uff1a")
    test_items = [
        "\u5355\u5143\uff1atest_ring_buffer, test_shm_layout, test_dma_manager, test_trace_log, test_eth_shm, test_link_model",
        "\u96c6\u6210 PCIe\uff1atest_sock_sync, test_bridge_loopback, test_dma_roundtrip, test_msi_roundtrip, test_precise_mode",
        "\u96c6\u6210 ETH\uff1atest_eth_loopback, test_link_drop, test_mac_stub_e2e, test_time_sync_loose",
        "\u5de5\u5177\uff1atest_cli_smoke, test_launch_smoke",
    ]
    for item in test_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7.3 VIP \u5192\u70df\u6d4b\u8bd5\uff08P5 \u65b0\u589e\uff09", level=2)
    add_para(doc, "\u6d4b\u8bd5 QEMU\u2194VCS \u5168\u94fe\u8def TLP \u5f80\u8fd4\uff08CfgRd0 completion + MWr + shutdown\uff09\uff1a")
    add_code_block(doc, """\
# \u7f16\u8bd1 QEMU \u4fa7\u6a21\u62df\u5668
gcc -o build/test_vip_smoke tests/e2e/test_vip_smoke.c \\
    bridge/qemu/bridge_qemu.c bridge/qemu/sock_sync.c \\
    bridge/common/shm_layout.c bridge/common/ring_buffer.c \\
    bridge/common/dma_manager.c bridge/common/trace_log.c \\
    -I bridge/common -I bridge/qemu -D_GNU_SOURCE -lrt -lpthread -std=c99

# \u7ec8\u7aef A: \u542f\u52a8 QEMU \u4fa7
./build/test_vip_smoke &

# \u7ec8\u7aef B: \u542f\u52a8 VCS
./build/simv_vip +SHM_NAME=/cosim_smoke +SOCK_PATH=/tmp/cosim_smoke.sock +UVM_TESTNAME=cosim_test""")

    # ========== 8. ETH \u6027\u80fd\u6d4b\u8bd5 ==========
    doc.add_heading("8. ETH \u6027\u80fd\u6d4b\u8bd5\uff08P5 \u65b0\u589e\uff09", level=1)

    doc.add_heading("8.1 \u5355\u5411\u541e\u5410\u91cf\u6d4b\u8bd5\uff08test_eth_iperf\uff09", level=2)
    add_para(doc, "iperf \u98ce\u683c\u7684 ETH SHM \u541e\u5410\u91cf\u57fa\u51c6\u6d4b\u8bd5\uff0c\u4f7f\u7528 fork \u521b\u5efa TX/RX \u4e24\u4e2a\u8fdb\u7a0b\u3002")
    add_para(doc, "\u7f16\u8bd1\uff1a", bold=True)
    add_code_block(doc, """\
gcc -o build/test_eth_iperf tests/e2e/test_eth_iperf.c \\
    bridge/common/eth_shm.c bridge/common/link_model.c \\
    bridge/eth/eth_port.c \\
    -I bridge/common -I bridge/eth -D_GNU_SOURCE -lrt -lpthread -std=c99 -O2""")
    add_para(doc, "\u7528\u6cd5\uff1a", bold=True)
    add_code_block(doc, """\
./build/test_eth_iperf [OPTIONS]
  -s SIZE    \u5e27\u8d1f\u8f7d\u5927\u5c0f\uff08\u5b57\u8282\uff0c\u9ed8\u8ba4 1500\uff09
  -t DUR     \u6d4b\u8bd5\u65f6\u957f\uff08\u79d2\uff0c\u9ed8\u8ba4 5\uff09
  -i INTV    \u62a5\u544a\u95f4\u9694\uff08\u79d2\uff0c\u9ed8\u8ba4 1\uff09
  -r RATE    \u9650\u901f Mbps\uff080=\u4e0d\u9650\uff0c\u9ed8\u8ba4 0\uff09
  -d DROP    \u4e22\u5305\u7387 ppm\uff08\u9ed8\u8ba4 0\uff09
  -l LAT     \u5355\u5411\u5ef6\u8fdf us\uff08\u9ed8\u8ba4 0\uff09
  -w WIN     \u6d41\u63a7\u7a97\u53e3\uff080=\u4e0d\u9650\uff0c\u9ed8\u8ba4 0\uff09""")
    add_para(doc, "\u793a\u4f8b\uff1a", bold=True)
    add_code_block(doc, """\
# 1500B \u6807\u51c6\u5e27\uff0c\u4e0d\u9650\u901f
./build/test_eth_iperf -s 1500 -t 5

# 9000B jumbo \u5e27
./build/test_eth_iperf -s 9000 -t 5

# \u5e26\u94fe\u8def\u6a21\u578b\uff1a10Gbps + 500ppm \u4e22\u5305 + 100us \u5ef6\u8fdf
./build/test_eth_iperf -s 1500 -t 5 -r 10000 -d 500 -l 100""")

    doc.add_heading("8.2 \u53cc\u5411\u4e92\u6253\u6d4b\u8bd5\uff08test_eth_bidir\uff09", level=2)
    add_para(doc, "\u6a21\u62df\u4e24\u4e2a QEMU+VCS \u8282\u70b9\u901a\u8fc7 ETH SHM \u53cc\u5411\u540c\u65f6\u901a\u4fe1\uff0c\u6bcf\u4e2a\u8fdb\u7a0b\u5185 TX/RX \u7ebf\u7a0b\u540c\u65f6\u5de5\u4f5c\u3002")
    add_para(doc, "\u7f16\u8bd1\uff1a", bold=True)
    add_code_block(doc, """\
gcc -o build/test_eth_bidir tests/e2e/test_eth_bidir.c \\
    bridge/common/eth_shm.c bridge/common/link_model.c \\
    bridge/eth/eth_port.c \\
    -I bridge/common -I bridge/eth -D_GNU_SOURCE -lrt -lpthread -std=c99 -O2""")
    add_para(doc, "\u7528\u6cd5\uff08\u53c2\u6570\u4e0e test_eth_iperf \u76f8\u540c\uff0c\u53bb\u6389 -w\uff09\uff1a", bold=True)
    add_code_block(doc, "./build/test_eth_bidir -s 1500 -t 5")

    doc.add_heading("8.3 \u6027\u80fd\u53c2\u8003\u6570\u636e", level=2)
    add_para(doc, "DEPTH=256, SHM=4.5MB", bold=True)
    add_para(doc, "\u5355\u5411\u541e\u5410\u91cf\uff1a", bold=True)
    make_table(doc,
        ["\u5e27\u5927\u5c0f", "\u541e\u5410\u91cf", "\u5e27\u7387", "\u5ef6\u8fdf avg"],
        [
            ["64B", "231 Mbps", "451 Kpps", "86 us"],
            ["1500B", "5,580 Mbps", "465 Kpps", "84 us"],
            ["9000B", "27,807 Mbps", "386 Kpps", "85 us"],
        ])
    doc.add_paragraph()
    add_para(doc, "\u53cc\u5411\u541e\u5410\u91cf\uff1a", bold=True)
    make_table(doc,
        ["\u5e27\u5927\u5c0f", "\u5355\u65b9\u5411 TX", "\u5355\u65b9\u5411 RX", "\u53cc\u5411\u603b\u8ba1", "\u5ef6\u8fdf avg"],
        [
            ["64B", "80 Mbps", "80 Mbps", "160 Mbps", "103 us"],
            ["1500B", "2,136 Mbps", "2,156 Mbps", "4,292 Mbps", "101 us"],
            ["9000B", "15,548 Mbps", "14,726 Mbps", "30,274 Mbps", "97 us"],
        ])

    doc.add_heading("8.4 Ring Depth \u8c03\u4f18", level=2)
    add_para(doc, "ETH_FRAME_RING_DEPTH \u9ed8\u8ba4 256\uff08\u6bcf\u65b9\u5411\uff09\uff0cSHM \u7ea6 4.5 MB\u3002\u53ef\u901a\u8fc7\u7f16\u8bd1\u53c2\u6570\u8986\u76d6\uff1a")
    add_code_block(doc, "gcc -DETH_FRAME_RING_DEPTH=128u ...")
    doc.add_paragraph()
    add_para(doc, "Ring Depth \u5bf9\u6bd4\uff081500B \u5e27\uff09\uff1a", bold=True)
    make_table(doc,
        ["Depth", "SHM \u5185\u5b58", "\u541e\u5410\u91cf", "\u5e27\u7387", "Send fails"],
        [
            ["64", "1.12 MB", "3,795 Mbps", "316 Kpps", "9,769"],
            ["128", "2.25 MB", "4,576 Mbps", "381 Kpps", "2"],
            ["256", "4.50 MB", "5,580 Mbps", "465 Kpps", "0"],
            ["512", "9.01 MB", "4,716 Mbps", "393 Kpps", "0"],
        ])
    add_para(doc, "\u5efa\u8bae\u4f7f\u7528\u9ed8\u8ba4\u503c 256\uff0c\u541e\u5410\u6700\u9ad8\u4e14 send fails \u4e3a\u96f6\u3002")

    # ========== 9. \u8fd0\u884c\u6a21\u5f0f ==========
    doc.add_heading("9. \u8fd0\u884c\u6a21\u5f0f", level=1)

    doc.add_heading("9.1 \u5feb\u901f\u6a21\u5f0f vs \u7cbe\u786e\u6a21\u5f0f", level=2)
    make_table(doc,
        ["\u6a21\u5f0f", "\u540c\u6b65\u7c92\u5ea6", "\u6027\u80fd", "\u9002\u7528\u573a\u666f"],
        [
            ["\u5feb\u901f\u6a21\u5f0f\uff08\u9ed8\u8ba4\uff09", "\u6bcf\u4e2a PCIe \u4e8b\u52a1", "~1000-10000 \u4e8b\u52a1/\u79d2", "\u9a71\u52a8\u52a0\u8f7d\u3001\u529f\u80fd\u9a8c\u8bc1"],
            ["\u7cbe\u786e\u6a21\u5f0f", "\u6bcf N \u4e2a\u65f6\u949f\u5468\u671f", "~10-100 \u4e8b\u52a1/\u79d2", "\u65f6\u5e8f\u8c03\u8bd5\u3001\u6ce2\u5f62\u5206\u6790"],
        ])

    doc.add_heading("9.2 \u8fd0\u884c\u65f6\u5207\u6362", level=2)
    add_code_block(doc, """\
#include "bridge_qemu.h"
bridge_request_mode_switch(ctx, COSIM_MODE_PRECISE);
bridge_advance_clock(ctx, 1000);
bridge_request_mode_switch(ctx, COSIM_MODE_FAST);""")

    doc.add_heading("9.3 \u4e8b\u52a1\u8ffd\u8e2a", level=2)
    add_code_block(doc, """\
bridge_enable_trace(ctx, "/tmp/cosim_trace.csv", TRACE_FMT_CSV);
// or JSON:
bridge_enable_trace(ctx, "/tmp/cosim_trace.json", TRACE_FMT_JSON);
bridge_disable_trace(ctx);""")

    # ========== 10. API \u53c2\u8003 ==========
    doc.add_heading("10. API \u53c2\u8003", level=1)

    doc.add_heading("10.1 QEMU \u4fa7 C API\uff08bridge_qemu.h\uff09", level=2)
    make_table(doc,
        ["\u51fd\u6570", "\u8bf4\u660e"],
        [
            ["bridge_init(shm_name, sock_path)", "\u521b\u5efa SHM \u5e76\u76d1\u542c socket"],
            ["bridge_connect(ctx)", "accept() VCS \u8fde\u63a5"],
            ["bridge_send_tlp(ctx, req)", "\u6392\u961f TLP \u8bf7\u6c42"],
            ["bridge_wait_completion(ctx, tag, cpl)", "\u7b49\u5f85 completion"],
            ["bridge_send_tlp_and_wait(ctx, req, cpl)", "\u53d1\u9001+\u7b49\u5f85\u5c01\u88c5"],
            ["bridge_send_tlp_fire(ctx, req)", "\u53d1\u9001 fire-and-forget TLP"],
            ["bridge_complete_dma(ctx, tag, status)", "DMA \u8bf7\u6c42\u5e94\u7b54"],
            ["bridge_request_mode_switch(ctx, mode)", "\u5207\u6362\u6a21\u5f0f"],
            ["bridge_advance_clock(ctx, cycles)", "\u7cbe\u786e\u6a21\u5f0f\u63a8\u8fdb\u5468\u671f"],
            ["bridge_enable_trace(ctx, path, fmt)", "\u5f00\u542f\u8ffd\u8e2a"],
            ["bridge_disable_trace(ctx)", "\u5173\u95ed\u8ffd\u8e2a"],
            ["bridge_destroy(ctx)", "\u91ca\u653e\u8d44\u6e90"],
        ])

    doc.add_heading("10.2 VCS \u4fa7 DPI-C \u51fd\u6570\uff08bridge_vcs.sv\uff09", level=2)
    add_code_block(doc, """\
import "DPI-C" function int  vcs_bridge_init_dpi(input string shm_name, input string sock_path);
import "DPI-C" function int  vcs_bridge_poll_tlp_dpi(...);
import "DPI-C" function int  vcs_bridge_send_completion_dpi(...);
import "DPI-C" function int  vcs_bridge_trigger_dma_dpi(...);
import "DPI-C" function int  vcs_bridge_raise_msi_dpi(input int vector);
import "DPI-C" function int  vcs_bridge_clock_ack_dpi();
import "DPI-C" function void vcs_bridge_close_dpi();""")

    doc.add_heading("10.3 ETH API\uff08eth_port.h\uff09", level=2)
    make_table(doc,
        ["\u51fd\u6570", "\u8bf4\u660e"],
        [
            ["eth_port_open(port, name, role, create)", "\u6253\u5f00 ETH \u7aef\u53e3"],
            ["eth_port_send(port, frame, now_ns)", "\u53d1\u5e27\uff080=OK, -1=full, -2=FC, -3=drop\uff09"],
            ["eth_port_recv(port, out, timeout_ns)", "\u6536\u5e27"],
            ["eth_port_tx_complete(port)", "FC \u786e\u8ba4"],
            ["eth_port_close(port)", "\u5173\u95ed\u7aef\u53e3"],
        ])

    doc.add_heading("10.4 \u94fe\u8def\u6a21\u578b\u53c2\u6570\uff08link_model.h\uff09", level=2)
    make_table(doc,
        ["\u5b57\u6bb5", "\u7c7b\u578b", "\u793a\u4f8b", "\u8bf4\u660e"],
        [
            ["drop_rate_ppm", "u32", "100000", "\u6bcf\u767e\u4e07\u5e27\u4e22\u5305\u6982\u7387\uff0810%\uff09"],
            ["burst_drop_len", "u16", "5", "\u89e6\u53d1\u540e\u8fde\u7eed\u4e22\u5e27\u6570"],
            ["latency_ns", "u64", "5000", "\u5355\u5411\u56fa\u5b9a\u5ef6\u8fdf\uff085us\uff09"],
            ["rate_mbps", "u32", "1000", "\u7ebf\u901f\u7387\uff0c0=\u65e0\u9650"],
            ["fc_window", "u32", "4", "\u6700\u5927\u5728\u98de\u5e27\u6570\uff0c0=\u65e0\u9650"],
        ])

    # ========== 11. \u53cc\u8282\u70b9 ETH \u4e92\u6253 ==========
    doc.add_heading("11. \u53cc\u8282\u70b9 ETH \u4e92\u6253", level=1)

    doc.add_heading("11.1 \u53cc\u8282\u70b9\u67b6\u6784", level=2)
    nodes = [
        "Node A: QEMU-A + VCS-A\uff0c\u901a\u8fc7 PCIe SHM \u4e92\u8054",
        "Node B: QEMU-B + VCS-B\uff0c\u901a\u8fc7 PCIe SHM \u4e92\u8054",
        "ETH SHM: A \u4e0e B \u8282\u70b9\u4e4b\u95f4\u7684\u53cc\u5411\u5e27\u961f\u5217",
    ]
    for n in nodes:
        doc.add_paragraph(n, style="List Bullet")

    doc.add_heading("11.2 \u542f\u52a8\u53cc\u8282\u70b9", level=2)
    add_code_block(doc, """\
python3 scripts/launch_dual.py \\
    --shm-pcie-a /cosim-pcie-a --sock-a /tmp/cosim-a.sock \\
    --shm-pcie-b /cosim-pcie-b --sock-b /tmp/cosim-b.sock \\
    --shm-eth   /cosim-eth0""")

    doc.add_heading("11.3 TAP Bridge \u6a21\u5f0f", level=2)
    add_code_block(doc, "./cosim.sh test tap")

    # ========== 12. VIP \u6a21\u5f0f\u8be6\u89e3 ==========
    doc.add_heading("12. VIP \u6a21\u5f0f\u8be6\u89e3\uff08P5 \u65b0\u589e\uff09", level=1)

    doc.add_heading("12.1 VIP \u6a21\u5f0f\u67b6\u6784", level=2)
    add_para(doc, "VIP \u6a21\u5f0f\u4f7f\u7528 UVM \u9a8c\u8bc1\u65b9\u6cd5\u5b66\uff0c\u901a\u8fc7 PCIe TL VIP \u66ff\u4ee3 Legacy \u6a21\u5f0f\u7684\u76f4\u63a5 DPI-C \u8c03\u7528\uff1a")
    vip_items = [
        "cosim_vip_top.sv: VIP \u6a21\u5f0f\u9876\u5c42\uff0c\u4f8b\u5316 pcie_tl_if\uff08\u8bf7\u6c42\u901a\u9053 + completion \u901a\u9053\uff09\u3001glue_if_to_stub\u3001pcie_ep_stub",
        "cosim_rc_driver: \u7ee7\u627f pcie_tl_rc_driver\uff0c\u901a\u8fc7 UVM factory override \u6ce8\u5165 DPI-C bridge \u903b\u8f91",
        "Completion \u901a\u8fc7\u7b2c\u4e8c\u4e2a pcie_tl_if (cpl_if) \u56de\u4f20\uff0c\u7ecf cosim_rc_driver \u89e3\u6790\u540e\u8f6c\u53d1\u7ed9 QEMU",
        "glue_if_to_stub: \u5c06 VIP 256-bit \u603b\u7ebf\u4fe1\u53f7\u8f6c\u6362\u4e3a pcie_ep_stub \u7684\u7b80\u5355 TLP \u4fe1\u53f7",
    ]
    for item in vip_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("12.2 UVM \u7ec4\u4ef6\u5173\u7cfb", level=2)
    add_code_block(doc, """\
cosim_test
  \u2514\u2500\u2500 cosim_env
       \u2514\u2500\u2500 rc_agent (pcie_tl_rc_agent, factory override driver \u2192 cosim_rc_driver)
            \u251c\u2500\u2500 cosim_rc_driver.request_loop()   \u2192 DPI-C poll TLP \u2192 drive VIP
            \u251c\u2500\u2500 cosim_rc_driver.completion_loop() \u2192 monitor cpl_if \u2192 handle_completion \u2192 forward to QEMU
            \u2514\u2500\u2500 cosim_rc_driver.dma_msi_loop()    \u2192 DMA/MSI \u2192 DPI-C""")

    doc.add_heading("12.3 \u7f16\u8bd1\u4e0e\u8fd0\u884c", level=2)
    add_code_block(doc, """\
make vcs-vip
./build/simv_vip +SHM_NAME=/cosim0 +SOCK_PATH=/tmp/cosim.sock +UVM_TESTNAME=cosim_test""")

    # ========== 13. DPU \u8fc1\u79fb\u96c6\u6210 ==========
    doc.add_heading("13. DPU \u8fc1\u79fb\u96c6\u6210", level=1)

    doc.add_heading("13.1 \u65b9\u6848\u4e00\uff1a\u7eaf RTL \u96c6\u6210", level=2)
    add_para(doc, "\u66ff\u6362 pcie_ep_stub.sv \u4e3a dpu_adapter.sv\uff0c\u8fde\u63a5\u771f\u5b9e DPU PCIe EP RTL\u3002")

    doc.add_heading("13.2 \u65b9\u6848\u4e8c\uff1aRTL + \u9a71\u52a8\u96c6\u6210", level=2)
    add_para(doc, "\u5728\u65b9\u6848\u4e00\u57fa\u7840\u4e0a\uff0cGuest Linux \u52a0\u8f7d DPU \u4e13\u7528\u9a71\u52a8\u3002")
    add_para(doc, "\u8be6\u7ec6\u65b9\u6848\u8bf7\u53c2\u8003 CoSim_DPU_Migration_Integration.docx\u3002")

    # ========== 14. \u7edf\u4e00\u547d\u4ee4\u884c\u5de5\u5177 ==========
    doc.add_heading("14. \u7edf\u4e00\u547d\u4ee4\u884c\u5de5\u5177\uff08cosim.sh\uff09", level=1)
    make_table(doc,
        ["\u547d\u4ee4", "\u7528\u6cd5", "\u8bf4\u660e"],
        [
            ["test", "./cosim.sh test <phase>", "\u8fd0\u884c\u6307\u5b9a\u6d4b\u8bd5\u9636\u6bb5"],
            ["start", "./cosim.sh start <component>", "\u542f\u52a8\u5355\u4e2a\u7ec4\u4ef6"],
            ["status", "./cosim.sh status", "\u67e5\u770b\u8fd0\u884c\u72b6\u6001"],
            ["log", "./cosim.sh log <component>", "\u67e5\u770b\u7ec4\u4ef6\u65e5\u5fd7"],
            ["clean", "./cosim.sh clean", "\u6e05\u7406 SHM \u548c\u4e34\u65f6\u6587\u4ef6"],
            ["info", "./cosim.sh info", "\u663e\u793a\u7cfb\u7edf\u4fe1\u606f"],
            ["help", "./cosim.sh help", "\u663e\u793a\u5e2e\u52a9"],
        ])

    # ========== 15. \u8c03\u8bd5\u5de5\u5177 ==========
    doc.add_heading("15. \u8c03\u8bd5\u5de5\u5177", level=1)

    doc.add_heading("15.1 cosim_cli", level=2)
    add_code_block(doc, "python3 scripts/cosim_cli.py --shm /cosim0 --sock /tmp/cosim.sock")

    doc.add_heading("15.2 trace_analyzer", level=2)
    add_code_block(doc, "python3 scripts/trace_analyzer.py /tmp/trace.csv")

    doc.add_heading("15.3 GDB \u8c03\u8bd5", level=2)
    add_code_block(doc, """\
GDB=1 ./scripts/run_cosim.sh
gdb vmlinux -ex 'target remote :1234'""")

    # ========== 16. \u6545\u969c\u6392\u67e5 ==========
    doc.add_heading("16. \u6545\u969c\u6392\u67e5", level=1)
    make_table(doc,
        ["\u73b0\u8c61", "\u53ef\u80fd\u539f\u56e0", "Ubuntu \u6392\u67e5\u6b65\u9aa4"],
        [
            ["bridge_init \u5931\u8d25", "/dev/shm \u7a7a\u95f4\u4e0d\u8db3", "df -h /dev/shm; rm -f /dev/shm/cosim*"],
            ["bridge_connect \u5361\u4f4f", "VCS \u672a\u542f\u52a8\u6216 sock_path \u4e0d\u4e00\u81f4", "ls -l /tmp/cosim.sock; \u786e\u8ba4\u4e24\u4fa7\u8def\u5f84"],
            ["VCS \u8bb8\u53ef\u8bc1\u9519\u8bef", "SNPSLMD_LICENSE_FILE \u672a\u8bbe\u7f6e", "export SNPSLMD_LICENSE_FILE=/path/to/license.dat"],
            ["QEMU \u7f16\u8bd1\u5931\u8d25", "\u7f3a\u5c11\u4f9d\u8d56", "sudo apt install libglib2.0-dev libpixman-1-dev libslirp-dev"],
            ["VCS \u7f16\u8bd1 C \u6587\u4ef6\u8def\u5f84\u9519\u8bef", "VCS \u5728 csrc/ \u5b50\u76ee\u5f55\u7f16\u8bd1", "\u4f7f\u7528 $(CURDIR) \u7edd\u5bf9\u8def\u5f84\uff08\u5df2\u5728 Makefile \u4e2d\u4fee\u590d\uff09"],
            ["DPI-C segfault", "fprintf/fflush \u5728 VCS \u9ad8\u9891\u8c03\u7528", "\u5220\u9664 poll \u5faa\u73af\u4e2d\u7684\u8c03\u8bd5\u6253\u5370"],
            ["Guest \u7f51\u7edc\u4e0d\u901a", "IP/MAC/ARP \u914d\u7f6e", "\u68c0\u67e5 config.env \u4e2d\u7684 Guest IP/MAC"],
            ["TAP \u8bbe\u5907\u521b\u5efa\u5931\u8d25", "\u6743\u9650\u4e0d\u8db3", "sudo chmod 666 /dev/net/tun \u6216\u4f7f\u7528 sudo"],
            ["Completion \u8d85\u65f6", "completion \u672a\u901a\u8fc7 VIP \u56de\u4f20", "\u68c0\u67e5 cpl_if \u8fde\u63a5\u548c cosim_rc_driver.completion_loop"],
        ])

    # ========== 17. \u9644\u5f55 ==========
    doc.add_heading("17. \u9644\u5f55", level=1)

    doc.add_heading("17.1 \u5171\u4eab\u5185\u5b58\u5e03\u5c40", level=2)
    make_table(doc,
        ["\u504f\u79fb", "\u5927\u5c0f", "\u5185\u5bb9"],
        [
            ["0x0000_0000", "4 KB", "\u63a7\u5236\u533a"],
            ["0x0000_1000", "256 KB", "\u8bf7\u6c42\u961f\u5217"],
            ["0x0004_1000", "256 KB", "\u54cd\u5e94\u961f\u5217"],
            ["0x0008_1000", "64 KB", "DMA + MSI \u961f\u5217"],
            ["0x0009_1000", "~63 MB", "DMA \u6570\u636e\u7f13\u51b2\u533a"],
        ])

    doc.add_heading("17.2 PCI ID", level=2)
    pci_items = [
        "Vendor ID: 0x1234",
        "Device ID: 0x0001",
        "Class: 0x0200 (Ethernet Controller)",
    ]
    for item in pci_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("17.3 \u6d88\u606f\u7c7b\u578b", level=2)
    msg_types = ("TLP_MWR, TLP_MRD, TLP_CFGWR, TLP_CFGRD0, TLP_CPL, "
                 "SYNC_MSG_TLP_READY, SYNC_MSG_CPL_READY, SYNC_MSG_DMA_REQ, "
                 "SYNC_MSG_DMA_CPL, SYNC_MSG_MSI, SYNC_MSG_MODE_SWITCH, "
                 "SYNC_MSG_CLOCK_STEP, SYNC_MSG_CLOCK_ACK, SYNC_MSG_SHUTDOWN")
    add_para(doc, msg_types)

    doc.add_heading("17.4 \u672f\u8bed\u8868", level=2)
    make_table(doc,
        ["\u7f29\u5199", "\u542b\u4e49"],
        [
            ["TLP", "Transaction Layer Packet"],
            ["MMIO", "Memory-Mapped I/O"],
            ["DMA", "Direct Memory Access"],
            ["MSI", "Message Signaled Interrupts"],
            ["SHM", "POSIX Shared Memory"],
            ["DPI-C", "Direct Programming Interface for C"],
            ["BAR", "Base Address Register"],
            ["EP", "Endpoint"],
            ["RC", "Root Complex"],
            ["VIP", "Verification IP"],
            ["UVM", "Universal Verification Methodology"],
            ["FC", "Flow Control"],
        ])

    doc.add_heading("17.5 \u7248\u672c\u5386\u53f2", level=2)
    make_table(doc,
        ["\u9636\u6bb5", "\u4ea4\u4ed8\u5185\u5bb9", "\u72b6\u6001"],
        [
            ["P1", "\u5355\u8282\u70b9 PCIe MMIO \u901a\u8def", "\u5df2\u5b8c\u6210"],
            ["P2", "DMA + MSI + \u7cbe\u786e\u6a21\u5f0f + Trace", "\u5df2\u5b8c\u6210"],
            ["P3", "\u53cc\u8282\u70b9 ETH \u4e92\u6253 + \u94fe\u8def\u6a21\u578b", "\u5df2\u5b8c\u6210"],
            ["P4", "cosim_cli, trace_analyzer, CI", "\u5df2\u5b8c\u6210"],
            ["P5", "UVM VIP \u6a21\u5f0f + ETH \u6027\u80fd\u57fa\u51c6", "\u5df2\u5b8c\u6210"],
        ])

    # Save
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    generate()
