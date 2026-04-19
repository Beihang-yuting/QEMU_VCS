#!/usr/bin/env python3
"""
生成 CoSim Platform DPU 迁移集成方案 Word 文档
包含两种方案:
  方案一: 纯 RTL 集成（替换 pcie_ep_stub.sv）
  方案二: RTL + DPU 软件驱动集成
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime


def add_code_block(doc, code):
    """添加代码块（灰色背景）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F0F0',
        qn('w:val'): 'clear'
    })
    pPr.append(shd)


def add_table(doc, headers, rows):
    """添加带表头的表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
    return table


def add_checklist(doc, items):
    """添加检查清单"""
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run(f'[ ] {item}')
        run.font.size = Pt(10)


def add_note(doc, text, note_type='注意'):
    """添加注意/提示框"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'[{note_type}] ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'FFF3E0',
        qn('w:val'): 'clear'
    })
    pPr.append(shd)


def create_document():
    doc = Document()

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ==================== 封面 ====================
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('QEMU-VCS CoSim Platform\nDPU 迁移集成方案')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xc4)

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        '方案一：纯 RTL 集成\n'
        '方案二：RTL + DPU 软件驱动集成'
    )
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f'版本: 1.0\n日期: {datetime.date.today().strftime("%Y-%m-%d")}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ==================== 目录 ====================
    doc.add_heading('目录', level=1)
    toc_items = [
        '1. 文档概述',
        '  1.1 文档目的',
        '  1.2 适用范围',
        '  1.3 术语定义',
        '2. CoSim 平台架构回顾',
        '  2.1 整体架构',
        '  2.2 可替换边界分析',
        '  2.3 TLP 接口定义',
        '3. 方案一：纯 RTL 集成',
        '  3.1 方案概述',
        '  3.2 前置条件',
        '  3.3 部署流程（6 步）',
        '  3.4 测试流程（4 阶段）',
        '  3.5 常见问题与排查',
        '4. 方案二：RTL + DPU 软件驱动集成',
        '  4.1 方案概述',
        '  4.2 与方案一的差异',
        '  4.3 前置条件',
        '  4.4 部署流程（8 步）',
        '  4.5 测试流程（6 阶段）',
        '  4.6 驱动调试技巧',
        '  4.7 常见问题与排查',
        '5. 两方案对比与选型建议',
        '6. 文件修改清单总览',
        '7. 附录',
        '  7.1 完整接口信号列表',
        '  7.2 DPI-C 函数参考',
        '  7.3 SHM 内存布局',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(10)
    doc.add_page_break()

    # ==================== 1. 文档概述 ====================
    doc.add_heading('1. 文档概述', level=1)

    doc.add_heading('1.1 文档目的', level=2)
    doc.add_paragraph(
        '本文档详细介绍如何将 QEMU-VCS CoSim 协同仿真平台迁移到已有的 DPU VCS 验证环境中。'
        '针对不同的集成深度，提供两种迁移方案：'
    )
    doc.add_paragraph(
        '方案一（纯 RTL 集成）：仅替换 RTL 侧的 PCIe EP Stub，使用裸寄存器访问和 '
        'DMA 测试工具验证硬件行为，适用于硬件验证阶段早期。',
        style='List Bullet'
    )
    doc.add_paragraph(
        '方案二（RTL + 驱动集成）：在方案一基础上，将 DPU 的完整软件驱动集成到 Guest Linux 中，'
        '通过驱动提供的标准接口（如 net_device、block_device）进行端到端测试，'
        '适用于软硬件协同验证阶段。',
        style='List Bullet'
    )

    doc.add_heading('1.2 适用范围', level=2)
    doc.add_paragraph(
        '本文档适用于：已有基于 VCS 的 DPU 验证平台，且 DPU RTL 已经具备 PCIe EP 功能'
        '（或 AXI/AXI-Stream 到 PCIe 桥接能力）的团队。CoSim 平台已在独立环境中验证通过'
        '（setup.sh 构建成功，所有核心测试 PASS）。'
    )

    doc.add_heading('1.3 术语定义', level=2)
    add_table(doc,
        ['术语', '说明'],
        [
            ['CoSim', 'Co-Simulation，QEMU 与 VCS 协同仿真'],
            ['TLP', 'Transaction Layer Packet，PCIe 事务层包'],
            ['SHM', 'Shared Memory，POSIX 共享内存'],
            ['DPI-C', 'Direct Programming Interface for C，SystemVerilog 调用 C 函数的标准接口'],
            ['EP', 'Endpoint，PCIe 端点设备（此处为 DPU）'],
            ['RC', 'Root Complex，PCIe 根复合体（QEMU 侧模拟）'],
            ['BAR', 'Base Address Register，PCI 设备基地址寄存器'],
            ['MSI', 'Message Signaled Interrupt，消息信号中断'],
            ['VQ', 'Virtqueue，VirtIO 虚拟队列'],
            ['CDC', 'Clock Domain Crossing，跨时钟域'],
            ['Initramfs', 'Initial RAM Filesystem，Linux 初始内存文件系统'],
        ]
    )

    doc.add_page_break()

    # ==================== 2. CoSim 平台架构回顾 ====================
    doc.add_heading('2. CoSim 平台架构回顾', level=1)

    doc.add_heading('2.1 整体架构', level=2)
    doc.add_paragraph(
        'CoSim 平台由以下核心组件组成：'
    )
    add_code_block(doc,
        '+--------------------------------------------------+\n'
        '|  Guest Linux (QEMU 虚拟机内)                      |\n'
        '|    测试程序 / DPU 驱动                             |\n'
        '+--------------------------------------------------+\n'
        '|  QEMU + cosim_pcie_rc 设备                       |\n'
        '|  (PCIe Root Complex 模拟)                        |\n'
        '+------------- SHM (共享内存) ----------------------+\n'
        '|  VCS Testbench                                   |\n'
        '|    tb_top.sv -> pcie_ep_stub.sv (可替换)          |\n'
        '|    bridge_vcs (DPI-C 桥接)                       |\n'
        '|    ETH MAC DPI (以太网转发)                       |\n'
        '+--------------------------------------------------+'
    )

    doc.add_paragraph(
        '各组件通过 POSIX 共享内存（4MB）通信。SHM 内部包含控制区（ctrl）、'
        '请求环形缓冲区（req_ring）、完成环形缓冲区（cpl_ring）、DMA 队列、'
        'MSI 队列和 DMA 数据缓冲区。'
    )

    doc.add_heading('2.2 可替换边界分析', level=2)
    doc.add_paragraph(
        'CoSim 平台的设计天然支持 DPU RTL 替换。可替换边界位于 tb_top.sv 与 '
        'pcie_ep_stub.sv 之间的 TLP 接口。这是一组标准化的信号，包含请求方向'
        '（Host -> EP）和完成方向（EP -> Host）两条通路。'
    )

    add_code_block(doc,
        '替换边界示意：\n\n'
        '  tb_top.sv (保留)\n'
        '      |\n'
        '      +-- DPI-C Bridge (bridge_vcs.c)     <-- 保留不动\n'
        '      |\n'
        '      +-- TLP 接口 ----------------------- <-- 替换点\n'
        '      |       |\n'
        '      |       +-- pcie_ep_stub.sv          <-- 原：VirtIO Stub\n'
        '      |       +-- dpu_adapter.sv           <-- 新：DPU RTL 适配\n'
        '      |\n'
        '      +-- ETH MAC DPI                     <-- 保留不动'
    )

    doc.add_heading('2.3 TLP 接口定义', level=2)
    doc.add_paragraph('以下是 CoSim 平台定义的 TLP 级接口信号，任何替换模块必须实现此接口：')

    doc.add_paragraph('请求方向（Host -> EP）：')
    add_table(doc,
        ['信号名', '方向', '位宽', '说明'],
        [
            ['tlp_valid', 'input', '1', 'TLP 请求有效标志'],
            ['tlp_type', 'input', '3', '请求类型：0=MRd, 1=MWr, 2=CfgRd, 3=CfgWr, 4=DMA_CPL'],
            ['tlp_addr', 'input', '64', '目标地址（BAR 偏移或 Config Space 地址）'],
            ['tlp_wdata', 'input', '32', '写数据（MWr/CfgWr 时有效）'],
            ['tlp_len', 'input', '16', '数据长度（字节）'],
            ['tlp_tag', 'input', '8', '事务标识符，用于匹配完成'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph('完成方向（EP -> Host）：')
    add_table(doc,
        ['信号名', '方向', '位宽', '说明'],
        [
            ['cpl_valid', 'output', '1', '完成有效标志'],
            ['cpl_tag', 'output', '8', '对应请求的 tag'],
            ['cpl_rdata', 'output', '32', '读返回数据（MRd/CfgRd 完成时有效）'],
            ['cpl_status', 'output', '1', '0=成功, 1=失败'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph('辅助信号（EP -> Host）：')
    add_table(doc,
        ['信号名', '方向', '位宽', '说明'],
        [
            ['notify_valid', 'output', '1', 'Virtio 通知有效（可选）'],
            ['notify_queue', 'output', '16', '被通知的队列编号（可选）'],
            ['isr_set', 'output', '1', '中断请求（触发 MSI）'],
        ]
    )

    doc.add_page_break()

    # ==================== 3. 方案一：纯 RTL 集成 ====================
    doc.add_heading('3. 方案一：纯 RTL 集成', level=1)

    doc.add_heading('3.1 方案概述', level=2)
    doc.add_paragraph(
        '本方案仅在 RTL 侧进行替换，将 pcie_ep_stub.sv 替换为真实的 DPU PCIe EP RTL。'
        'Guest Linux 内不加载任何 DPU 专用驱动，仅使用通用测试工具（devmem、cfgspace_test 等）'
        '直接访问寄存器和 DMA 通路，验证硬件功能的正确性。'
    )
    doc.add_paragraph('适用场景：')
    for s in [
        'DPU RTL 开发早期，验证 PCIe EP 基本功能',
        'Config Space / BAR 寄存器映射验证',
        'DMA 引擎功能验证',
        '中断路径验证',
        '软件驱动尚未就绪的情况',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('3.2 前置条件', level=2)
    add_table(doc,
        ['条件', '要求', '检查方法'],
        [
            ['CoSim 平台', '已通过 setup.sh 构建', 'ls build/libcosim_bridge.so'],
            ['QEMU', '已编译且包含 cosim-pcie-rc 设备', 'qemu-system-x86_64 --version'],
            ['VCS', '已安装，VCS_HOME 正确设置', 'which vcs'],
            ['DPU RTL', 'PCIe EP 模块可用', '确认源文件列表'],
            ['Guest 内核', 'vmlinuz + initramfs 可用', 'ls scripts/vmlinuz'],
        ]
    )

    doc.add_heading('3.3 部署流程（6 步）', level=2)

    # Step 1
    doc.add_heading('Step 1: 环境准备', level=3)
    doc.add_paragraph('在 DPU VCS 验证服务器上克隆并构建 CoSim 平台基础组件：')
    add_code_block(doc,
        '# 1. 克隆 CoSim 平台代码\n'
        'git clone <cosim-platform-repo> ~/cosim-platform\n'
        'cd ~/cosim-platform\n'
        '\n'
        '# 2. 运行 setup.sh 安装基础依赖\n'
        '#    GCC 4.8 环境已适配（compat_atomic.h）\n'
        '#    VCS_HOME 会自动从 PATH 推导\n'
        'bash setup.sh\n'
        '\n'
        '# 3. 确认构建产物\n'
        '#    必须成功: libcosim_bridge.so, qemu-system-x86_64\n'
        '#    simv 暂不构建（需要适配后重新编译）\n'
        'ls build/libcosim_bridge.so\n'
        'ls ~/workspace/qemu-9.2.0/build/qemu-system-x86_64'
    )

    # Step 2
    doc.add_heading('Step 2: 识别 DPU PCIe 接口类型', level=3)
    doc.add_paragraph(
        '确认 DPU RTL 暴露的 PCIe 接口类型，决定适配层的复杂度：'
    )
    add_table(doc,
        ['接口类型', '适配复杂度', '说明'],
        [
            ['TLP 级接口', '低 -- 直接信号映射', 'DPU 已有 tlp_valid/type/addr 等信号'],
            ['AXI-Stream PCIe', '中 -- 需 TLP<->AXI-S 桥', 'DPU 使用 AXI-Stream 封装 TLP'],
            ['AXI Memory-Mapped', '中高 -- 需协议转换', 'DPU 使用 AXI4/AXI4-Lite MMIO'],
            ['Xilinx PCIe IP', '中 -- 需适配 IP 接口', '使用 Xilinx 的 PCIe Gen3/4 IP'],
        ]
    )
    add_note(doc,
        '如果 DPU 使用标准 TLP 级接口，适配层仅为信号名映射，约 50 行代码。'
        '如果使用 AXI 接口，需要编写 TLP<->AXI 协议转换桥，约 200-500 行。'
    )

    # Step 3
    doc.add_heading('Step 3: 编写接口适配层 (dpu_adapter.sv)', level=3)
    doc.add_paragraph('创建 dpu_adapter.sv，替换 pcie_ep_stub.sv：')
    add_code_block(doc,
        '// dpu_adapter.sv - 将 CoSim TLP 接口连接到真实 DPU PCIe EP\n'
        'module dpu_adapter (\n'
        '    input  logic        clk,\n'
        '    input  logic        rst_n,\n'
        '\n'
        '    // CoSim 侧 TLP 接口 (来自 tb_top)\n'
        '    input  logic        tlp_valid,\n'
        '    input  logic [2:0]  tlp_type,\n'
        '    input  logic [63:0] tlp_addr,\n'
        '    input  logic [31:0] tlp_wdata,\n'
        '    input  logic [15:0] tlp_len,\n'
        '    input  logic [7:0]  tlp_tag,\n'
        '    output logic        cpl_valid,\n'
        '    output logic [7:0]  cpl_tag,\n'
        '    output logic [31:0] cpl_rdata,\n'
        '    output logic        cpl_status,\n'
        '    output logic        notify_valid,\n'
        '    output logic [15:0] notify_queue,\n'
        '    output logic        isr_set\n'
        ');\n'
        '\n'
        '    // ---- 情况 A: DPU 已有 TLP 接口，直接映射 ----\n'
        '    your_dpu_pcie_ep u_dpu_ep (\n'
        '        .clk           (clk),\n'
        '        .rst_n         (rst_n),\n'
        '        .pcie_req_vld  (tlp_valid),   // 信号名映射\n'
        '        .pcie_req_type (tlp_type),\n'
        '        .pcie_req_addr (tlp_addr),\n'
        '        .pcie_req_data (tlp_wdata),\n'
        '        .pcie_req_len  (tlp_len),\n'
        '        .pcie_req_tag  (tlp_tag),\n'
        '        .pcie_cpl_vld  (cpl_valid),\n'
        '        .pcie_cpl_tag  (cpl_tag),\n'
        '        .pcie_cpl_data (cpl_rdata),\n'
        '        .pcie_cpl_err  (cpl_status),\n'
        '        .pcie_intr     (isr_set)\n'
        '    );\n'
        '\n'
        '    // ---- 情况 B: DPU 用 AXI 接口 ----\n'
        '    // 需要在此实例化 TLP<->AXI 协议转换桥\n'
        '    // tlp_to_axi_bridge u_bridge (\n'
        '    //     .clk(clk), .rst_n(rst_n),\n'
        '    //     .tlp_*(...),  // CoSim 侧\n'
        '    //     .axi_*(...),  // DPU 侧\n'
        '    // );\n'
        '\n'
        '    // notify_valid / notify_queue 可选\n'
        '    // 如果 DPU 不使用 Virtio，可直接置 0\n'
        '    // assign notify_valid = 1\'b0;\n'
        '    // assign notify_queue = 16\'h0;\n'
        '\n'
        'endmodule'
    )

    # Step 4
    doc.add_heading('Step 4: 修改 tb_top.sv', level=3)
    doc.add_paragraph('将 pcie_ep_stub 实例替换为 dpu_adapter：')
    add_code_block(doc,
        '// tb_top.sv 中修改:\n'
        '\n'
        '// ---- 原始代码 ----\n'
        '// pcie_ep_stub u_ep (\n'
        '//     .clk(clk), .rst_n(rst_n),\n'
        '//     .tlp_valid(ep_tlp_valid), ...\n'
        '// );\n'
        '\n'
        '// ---- 替换为 ----\n'
        'dpu_adapter u_dpu (\n'
        '    .clk          (clk),\n'
        '    .rst_n        (rst_n),\n'
        '    .tlp_valid    (ep_tlp_valid),\n'
        '    .tlp_type     (ep_tlp_type),\n'
        '    .tlp_addr     (ep_tlp_addr),\n'
        '    .tlp_wdata    (ep_tlp_wdata),\n'
        '    .tlp_len      (ep_tlp_len),\n'
        '    .tlp_tag      (ep_tlp_tag),\n'
        '    .cpl_valid    (ep_cpl_valid),\n'
        '    .cpl_tag      (ep_cpl_tag),\n'
        '    .cpl_rdata    (ep_cpl_rdata),\n'
        '    .cpl_status   (ep_cpl_status),\n'
        '    .notify_valid (ep_notify_valid),\n'
        '    .notify_queue (ep_notify_queue),\n'
        '    .isr_set      (ep_isr_set)\n'
        ');\n'
        '\n'
        '// 注意: 如果 DPU 有独立时钟域,\n'
        '// 需要在 dpu_adapter 中处理 CDC 同步'
    )
    add_note(doc,
        '如果 DPU 内部有自己的时钟域（与 CoSim 的 clk 不同频），'
        '需要在 dpu_adapter.sv 中添加 CDC（跨时钟域同步）逻辑，'
        '包括异步 FIFO 或双触发器同步链。'
    )

    # Step 5
    doc.add_heading('Step 5: 编写 VCS 编译脚本', level=3)
    doc.add_paragraph('创建 build_dpu_simv.sh，将 DPU RTL 与 CoSim 框架一起编译：')
    add_code_block(doc,
        '#!/bin/bash\n'
        '# build_dpu_simv.sh - 编译 DPU 版 simv\n'
        '\n'
        'set -euo pipefail\n'
        'source ~/set-env.sh 2>/dev/null || true\n'
        '\n'
        '# ========== 配置区（根据实际项目修改）==========\n'
        '\n'
        '# DPU RTL 文件列表\n'
        'DPU_RTL_FILES=(\n'
        '    /path/to/dpu/rtl/pcie_ep_top.sv\n'
        '    /path/to/dpu/rtl/pcie_cfg_space.sv\n'
        '    /path/to/dpu/rtl/dma_engine.sv\n'
        '    /path/to/dpu/rtl/bar_decoder.sv\n'
        '    # ... 添加所有 DPU RTL 文件 ...\n'
        ')\n'
        '\n'
        '# DPU RTL Include 路径\n'
        'DPU_INCDIRS=(\n'
        '    +incdir+/path/to/dpu/rtl/includes\n'
        '    +incdir+/path/to/dpu/rtl/common\n'
        ')\n'
        '\n'
        '# ========== CoSim 框架文件（通常无需修改）==========\n'
        '\n'
        'COSIM_SV_FILES=(\n'
        '    bridge/vcs/bridge_vcs.sv\n'
        '    vcs-tb/dpu_adapter.sv       # 新的适配层\n'
        '    vcs-tb/tb_top.sv            # 修改后的 tb_top\n'
        ')\n'
        '\n'
        'COSIM_C_FILES=(\n'
        '    bridge/vcs/bridge_vcs.c\n'
        '    bridge/common/shm_layout.c\n'
        '    bridge/common/ring_buffer.c\n'
        '    bridge/common/dma_manager.c\n'
        '    bridge/common/trace_log.c\n'
        '    bridge/eth/eth_port.c\n'
        '    bridge/eth/eth_mac_dpi.c\n'
        '    bridge/vcs/virtqueue_dma.c\n'
        ')\n'
        '\n'
        '# ========== 编译 ==========\n'
        '\n'
        'echo "==== Building DPU simv ===="\n'
        'vcs -full64 -sverilog \\\n'
        '    -timescale=1ns/1ps \\\n'
        '    -CFLAGS "-std=gnu99 -I bridge/common -I bridge/eth" \\\n'
        '    -LDFLAGS "-lrt -lpthread" \\\n'
        '    ${COSIM_SV_FILES[@]} \\\n'
        '    ${DPU_RTL_FILES[@]} \\\n'
        '    ${COSIM_C_FILES[@]} \\\n'
        '    ${DPU_INCDIRS[@]} \\\n'
        '    +define+DPU_MODE \\\n'
        '    -o simv_dpu \\\n'
        '    -l compile.log\n'
        '\n'
        'echo "==== Build complete: ./simv_dpu ===="'
    )

    # Step 6
    doc.add_heading('Step 6: 启动验证', level=3)
    doc.add_paragraph('QEMU 侧无需任何修改，SHM 协议完全透明：')
    add_code_block(doc,
        '# 终端 1: 启动 VCS 仿真\n'
        './simv_dpu +cosim_shm=/cosim_shm &\n'
        '# 等待日志: "vcs_ready = 1"\n'
        '\n'
        '# 终端 2: 启动 QEMU\n'
        'qemu-system-x86_64 \\\n'
        '    -device cosim-pcie-rc,shm_name=/cosim_shm \\\n'
        '    -kernel scripts/vmlinuz \\\n'
        '    -initrd scripts/initramfs.cpio.gz \\\n'
        '    -append "console=ttyS0" \\\n'
        '    -nographic -m 512M\n'
        '\n'
        '# Guest 启动后, 运行测试工具:\n'
        './cfgspace_test        # 读取 Config Space\n'
        './devmem_test          # BAR0 寄存器读写\n'
        './dma_test             # DMA 通路验证'
    )

    doc.add_heading('3.4 测试流程（4 阶段）', level=2)

    # 阶段 1
    doc.add_heading('阶段 1: 接口握手验证', level=3)
    doc.add_paragraph('目标：确认 SHM 通信链路建立，DPU EP 能响应 TLP。')
    add_code_block(doc,
        '# 1. 观察 VCS 日志\n'
        '#    预期: "SHM opened", "vcs_ready = 1"\n'
        '\n'
        '# 2. 观察 QEMU 日志\n'
        '#    预期: "cosim: qemu_ready=1, waiting for vcs..."\n'
        '#    然后: "cosim: vcs ready, starting..."\n'
        '\n'
        '# 3. Guest 内运行\n'
        './cfgspace_test\n'
        '# 预期: 读到 DPU 的 Vendor/Device ID'
    )
    add_checklist(doc, [
        'SHM magic/version 匹配',
        'qemu_ready 和 vcs_ready 均为 1',
        'Config Space 读返回正确的 Vendor/Device ID',
        'Subsystem ID 和 Class Code 正确',
    ])

    # 阶段 2
    doc.add_heading('阶段 2: 寄存器访问验证', level=3)
    doc.add_paragraph('目标：验证 DPU BAR0 寄存器映射正确。')
    add_code_block(doc,
        '# Guest 内:\n'
        './devmem_test <bar0_base> <register_offset>\n'
        '\n'
        '# 批量验证所有寄存器:\n'
        'for offset in 0x00 0x04 0x08 0x0c 0x10; do\n'
        '    ./devmem_test $BAR0 $offset\n'
        'done'
    )
    add_checklist(doc, [
        '所有 BAR0 寄存器读写正确',
        '寄存器复位值符合设计规范',
        '只读寄存器写入被忽略（读回值不变）',
        '保留位读回为 0',
    ])

    # 阶段 3
    doc.add_heading('阶段 3: DMA 通路验证', level=3)
    doc.add_paragraph('目标：验证 DPU 发起的 DMA 读写能正确到达 Guest 物理内存。')
    add_code_block(doc,
        '# DMA 数据流:\n'
        '# DPU RTL -> dpu_adapter -> tb_top (handle_dma_doorbell)\n'
        '#   -> bridge_vcs_dma_read_sync / bridge_vcs_dma_write_sync\n'
        '#   -> SHM dma_req_ring -> QEMU irq_poller\n'
        '#   -> cpu_physical_memory_read/write\n'
        '#   -> SHM dma_cpl_ring -> VCS 侧返回\n'
        '\n'
        '# Guest 内:\n'
        './dma_test'
    )
    add_checklist(doc, [
        'DMA Read 返回正确数据',
        'DMA Write 数据正确写入 Guest 内存',
        'DMA 完成中断正确触发',
        '大块 DMA 传输（>4KB）正确完成',
    ])

    # 阶段 4
    doc.add_heading('阶段 4: 数据通路端到端验证', level=3)
    doc.add_paragraph('目标：验证完整的 Guest <-> DPU <-> Network 数据流（如果 DPU 支持网络功能）。')
    add_code_block(doc,
        '# 如果 DPU 实现了网络功能:\n'
        '\n'
        '# 终端 1: Host 侧启动 ETH 接收端\n'
        './eth_receiver\n'
        '\n'
        '# 终端 2: Guest 内发送测试帧\n'
        './nic_tx_test\n'
        '\n'
        '# 验证 eth_receiver 收到帧\n'
        '# 检查帧内容、长度、校验和'
    )
    add_checklist(doc, [
        'TX: Guest 写 VQ -> DPU 处理 -> ETH SHM -> 外部接收',
        'RX: 外部发送 -> ETH SHM -> DPU 处理 -> VQ -> Guest 接收',
        'MSI 中断正确送达 Guest',
    ])

    doc.add_heading('3.5 常见问题与排查', level=2)
    add_table(doc,
        ['现象', '原因', '解决方案'],
        [
            ['CfgRd 返回全 0xFFFFFFFF',
             'DPU EP 未连接或 tlp_type 编码不匹配',
             '检查 dpu_adapter 中 tlp_type 映射是否与 DPU 一致'],
            ['DMA 超时无返回',
             'DMA doorbell 未触发或 dma_cpl 未写回',
             '在 dpu_adapter 中加 $display 跟踪 DMA 请求'],
            ['仿真挂死不前进',
             '握手死锁：QEMU 等 VCS 完成，VCS 等 DMA 完成',
             '检查 irq_poller 线程是否正常处理 dma_req_ring'],
            ['VCS 编译报 undefined module',
             'DPU RTL 文件列表不完整',
             '检查所有依赖模块是否加入编译列表'],
            ['时序违例 / 竞态',
             'DPU 时钟域与 CoSim 不匹配',
             '在 dpu_adapter 中加 CDC 同步逻辑'],
            ['VCS 编译报 timescale 错误',
             '混合文件 timescale 不一致',
             '使用 -timescale=1ns/1ps 全局覆盖'],
        ]
    )

    doc.add_page_break()

    # ==================== 4. 方案二：RTL + 驱动集成 ====================
    doc.add_heading('4. 方案二：RTL + DPU 软件驱动集成', level=1)

    doc.add_heading('4.1 方案概述', level=2)
    doc.add_paragraph(
        '本方案在方案一（纯 RTL 集成）的基础上，进一步将 DPU 的完整软件驱动集成到 '
        'Guest Linux 中。驱动通过 Linux PCI 子系统 probe DPU 设备，并提供标准的内核接口'
        '（如 net_device 或 block_device），使得可以使用标准 Linux 工具（ip、ethtool、'
        'iperf3 等）进行端到端测试。'
    )
    doc.add_paragraph('适用场景：')
    for s in [
        '软硬件协同验证阶段，驱动已具备基本功能',
        '验证驱动与硬件的寄存器交互是否一致',
        '验证 DMA 引擎与驱动的 DMA 映射是否正确',
        '验证中断路径（MSI/MSI-X）从硬件到驱动的完整链路',
        '使用驱动原生测试工具进行功能回归',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('4.2 与方案一的差异', level=2)
    add_table(doc,
        ['维度', '方案一（纯 RTL）', '方案二（RTL + 驱动）'],
        [
            ['Guest 内测试方式', 'devmem/cfgspace_test 裸访问', '通过驱动标准接口（ip/ethtool/iperf）'],
            ['需要修改的文件', '3 个（adapter + tb_top + build 脚本）', '6+ 个（额外：QEMU 插件 + initramfs + init 脚本）'],
            ['PCI ID 匹配', '不需要（裸访问不依赖驱动 probe）', '必须三处一致（QEMU/RTL/驱动）'],
            ['Initramfs', '使用原始版本', '需要重新打包（含驱动 .ko）'],
            ['测试深度', '寄存器级 + DMA 级', '应用层端到端（TCP/UDP/RDMA）'],
            ['调试难度', '低（直接看波形）', '中（需要同时看波形 + 驱动日志）'],
            ['部署复杂度', '低', '中'],
        ]
    )

    doc.add_heading('4.3 前置条件', level=2)
    doc.add_paragraph('除方案一的所有前置条件外，还需要：')
    add_table(doc,
        ['条件', '要求', '检查方法'],
        [
            ['DPU 驱动源码', '可编译的内核模块（.ko）', 'ls dpu_driver/*.c Makefile/Kbuild'],
            ['Guest 内核头文件', '与 vmlinuz 版本完全匹配', 'uname -r vs kernel headers version'],
            ['驱动 PCI ID', '已知 Vendor ID + Device ID', 'grep PCI_DEVICE dpu_driver/*.c'],
            ['驱动依赖模块', '确认内核 CONFIG 选项', 'modinfo dpu_net.ko | grep depends'],
            ['交叉编译工具链', '如果 Host 与 Guest 架构不同', 'which x86_64-linux-gnu-gcc'],
        ]
    )

    doc.add_heading('4.4 部署流程（8 步）', level=2)

    doc.add_paragraph(
        '方案二的前 5 步与方案一完全相同（环境准备、接口识别、编写 adapter、修改 tb_top、编译 simv）。'
        '以下仅描述新增和修改的步骤。'
    )
    add_note(doc,
        '如果你已完成方案一的 Step 1-5，可直接从 Step 6 开始。'
    )

    # Step 6
    doc.add_heading('Step 6: PCI ID 三点对齐（关键步骤）', level=3)
    doc.add_paragraph(
        '这是方案二与方案一最关键的区别。驱动的 PCI probe 机制要求设备的 Vendor ID 和 '
        'Device ID 必须与驱动注册的 pci_device_id 完全匹配，否则驱动不会 probe。'
        '必须确保以下三处 ID 一致：'
    )
    add_table(doc,
        ['位置', '文件', '字段/信号', '如何修改'],
        [
            ['(1) QEMU 插件', 'qemu-plugin/cosim_pcie_rc.c',
             'pci_config_set_vendor_id()\npci_config_set_device_id()',
             '改为 DPU 驱动期望的 VID/DID'],
            ['(2) DPU RTL', 'Config Space 寄存器（地址 0x00）',
             'Vendor ID [15:0] + Device ID [31:16]',
             '确认 RTL 实现的值'],
            ['(3) DPU 驱动', 'dpu_driver.c 中的 pci_device_id 表',
             'PCI_DEVICE(vid, did)',
             '通常不改驱动，改 QEMU 和 RTL 去匹配'],
        ]
    )

    doc.add_paragraph('修改 QEMU 插件示例：')
    add_code_block(doc,
        '// cosim_pcie_rc.c 中修改:\n'
        '\n'
        '// 找到 PCI 配置初始化处，修改为 DPU 的 ID:\n'
        'pci_config_set_vendor_id(pci_conf, 0x1234);    // 改为 DPU VID\n'
        'pci_config_set_device_id(pci_conf, 0x5678);    // 改为 DPU DID\n'
        '\n'
        '// 同时确保以下字段也匹配:\n'
        'pci_config_set_class(pci_conf, PCI_CLASS_NETWORK_ETHERNET);  // 或 DPU 实际 Class\n'
        'pci_conf[PCI_REVISION_ID] = 0x01;              // Revision\n'
        'pci_config_set_subsystem_vendor_id(pci_conf, 0x1234);  // Subsystem VID\n'
        'pci_config_set_subsystem_id(pci_conf, 0x0001);         // Subsystem DID\n'
        '\n'
        '// BAR 大小也需要匹配驱动期望:\n'
        'memory_region_init_io(&s->bar0, OBJECT(s), &cosim_bar0_ops,\n'
        '                     s, "cosim-bar0", DPU_BAR0_SIZE);  // 改为实际 BAR 大小'
    )

    add_note(doc,
        '三处 ID 不一致是方案二最常见的失败原因。请在编译前反复核对！'
        '使用 "grep -r PCI_DEVICE dpu_driver/" 查找驱动期望的 ID。',
        '重要'
    )

    # Step 7
    doc.add_heading('Step 7: DPU 驱动编译与 Initramfs 打包', level=3)

    doc.add_paragraph('7.1 编译驱动')
    add_code_block(doc,
        '# 驱动需要针对 Guest 内核编译（不是 Host 内核）\n'
        '# KDIR 指向与 Guest vmlinuz 版本匹配的内核源码/头文件\n'
        '\n'
        'KDIR=/path/to/guest-kernel-source\n'
        '\n'
        'cd /path/to/dpu-driver\n'
        'make KDIR=$KDIR ARCH=x86_64 CROSS_COMPILE= modules\n'
        '\n'
        '# 产物:\n'
        'ls *.ko\n'
        '# 例如: dpu_net.ko\n'
        '\n'
        '# 检查驱动依赖:\n'
        'modinfo dpu_net.ko | grep depends\n'
        '# 如果有依赖（如 pci_hyperv, uio），也需要打入 initramfs'
    )

    add_note(doc,
        '内核版本必须完全匹配！如果编译驱动用的 KDIR 与 Guest vmlinuz 版本不同，'
        'insmod 会报 "version magic" 错误而拒绝加载。',
        '重要'
    )

    doc.add_paragraph('7.2 将驱动打入 Initramfs')
    doc.add_paragraph('提供两种方式：')

    doc.add_paragraph('方式 A：编译时内建（CONFIG_DPU_NET=y）')
    add_code_block(doc,
        '# 将驱动编译进内核（不作为模块）\n'
        '# 修改 Guest 内核 .config:\n'
        '#   CONFIG_DPU_NET=y  (而非 =m)\n'
        '#\n'
        '# 重新编译内核:\n'
        'cd /path/to/guest-kernel-source\n'
        'make -j$(nproc)\n'
        '\n'
        '# 使用新的 vmlinuz 启动 QEMU\n'
        '# 优点: 无需 insmod，驱动随内核启动自动可用\n'
        '# 缺点: 每次修改驱动需要重编内核'
    )

    doc.add_paragraph('方式 B：打包进 initramfs（推荐，更灵活）')
    add_code_block(doc,
        '#!/bin/bash\n'
        '# rebuild_initramfs_with_driver.sh\n'
        '\n'
        'ORIG_INITRAMFS=/path/to/initramfs.cpio.gz\n'
        'DRIVER_DIR=/path/to/dpu-driver\n'
        'WORK_DIR=/tmp/initramfs_work\n'
        'OUTPUT=/path/to/initramfs_with_driver.cpio.gz\n'
        '\n'
        '# 1. 解包现有 initramfs\n'
        'rm -rf $WORK_DIR && mkdir -p $WORK_DIR\n'
        'cd $WORK_DIR\n'
        'zcat $ORIG_INITRAMFS | cpio -idmv 2>/dev/null\n'
        '\n'
        '# 2. 放入驱动模块\n'
        'mkdir -p lib/modules\n'
        'cp $DRIVER_DIR/*.ko lib/modules/\n'
        'echo "Copied drivers: $(ls lib/modules/*.ko)"\n'
        '\n'
        '# 3. 放入固件文件（如果需要）\n'
        '# mkdir -p lib/firmware/dpu\n'
        '# cp /path/to/firmware/* lib/firmware/dpu/\n'
        '\n'
        '# 4. 放入驱动测试工具（如果有）\n'
        '# cp /path/to/dpu_test_tool usr/bin/\n'
        '\n'
        '# 5. 修改 init 脚本，添加驱动加载\n'
        'cat >> init << \'INIT_APPEND\'\n'
        '\n'
        '# ---- DPU 驱动加载 ----\n'
        'echo "[DPU] Loading DPU driver..."\n'
        'insmod /lib/modules/dpu_net.ko\n'
        'sleep 1\n'
        '\n'
        '# 检查驱动是否识别到设备\n'
        'if dmesg | grep -q "dpu.*probe"; then\n'
        '    echo "[DPU] Driver probe successful"\n'
        '    # 如果是网卡驱动，配置网络接口\n'
        '    if ip link show eth0 > /dev/null 2>&1; then\n'
        '        ip link set eth0 up\n'
        '        ip addr add 192.168.100.2/24 dev eth0\n'
        '        echo "[DPU] eth0 configured: 192.168.100.2/24"\n'
        '    fi\n'
        'else\n'
        '    echo "[DPU] WARNING: Driver loaded but no device found"\n'
        '    echo "[DPU] lspci output:"\n'
        '    lspci 2>/dev/null || echo "  lspci not available"\n'
        '    echo "[DPU] dmesg tail:"\n'
        '    dmesg | tail -20\n'
        'fi\n'
        'INIT_APPEND\n'
        '\n'
        '# 6. 重新打包\n'
        'find . | cpio -o -H newc 2>/dev/null | gzip > $OUTPUT\n'
        'echo "Output: $OUTPUT ($(du -h $OUTPUT | cut -f1))"\n'
        'echo "Done."'
    )

    # Step 8
    doc.add_heading('Step 8: 启动验证（含驱动）', level=3)
    doc.add_paragraph('使用包含驱动的 initramfs 启动 QEMU：')
    add_code_block(doc,
        '# 终端 1: 启动 VCS 仿真（与方案一相同）\n'
        './simv_dpu +cosim_shm=/cosim_shm &\n'
        '\n'
        '# 终端 2: 启动 QEMU（注意 initramfs 改为含驱动版本）\n'
        'qemu-system-x86_64 \\\n'
        '    -device cosim-pcie-rc,shm_name=/cosim_shm \\\n'
        '    -kernel scripts/vmlinuz \\\n'
        '    -initrd initramfs_with_driver.cpio.gz \\\n'
        '    -append "console=ttyS0 loglevel=7" \\\n'
        '    -nographic -m 512M\n'
        '\n'
        '# loglevel=7 开启最大日志，便于观察驱动 probe 过程'
    )

    doc.add_heading('4.5 测试流程（6 阶段）', level=2)
    doc.add_paragraph(
        '方案二的测试流程在方案一的 4 个阶段基础上，新增 2 个驱动相关阶段，'
        '共 6 个阶段。阶段 1-2 与方案一完全相同。'
    )

    # 阶段 1-2 引用方案一
    doc.add_heading('阶段 1-2: 接口握手 + 寄存器验证', level=3)
    doc.add_paragraph('与方案一的阶段 1-2 完全相同，参见第 3.4 节。确保基础硬件功能正常后再继续。')

    # 阶段 3: 驱动 Probe
    doc.add_heading('阶段 3: 驱动 Probe 验证（新增）', level=3)
    doc.add_paragraph('目标：确认驱动正确识别并绑定 DPU 设备。')
    add_code_block(doc,
        '# Guest 内:\n'
        '\n'
        '# 1. 确认 PCI 设备可见\n'
        'lspci\n'
        '# 预期: 能看到 DPU 设备（VID:DID 匹配）\n'
        '# 示例输出: 00:04.0 Ethernet controller: Vendor 1234 Device 5678\n'
        '\n'
        '# 2. 确认驱动已绑定\n'
        'lspci -k\n'
        '# 预期: "Kernel driver in use: dpu_net"\n'
        '\n'
        '# 3. 查看驱动日志\n'
        'dmesg | grep -i dpu\n'
        '# 预期: probe 成功，无错误\n'
        '# 示例: "dpu_net 0000:00:04.0: DPU NIC initialized"\n'
        '\n'
        '# 4. 查看 BAR 映射\n'
        'lspci -v\n'
        '# 预期: BAR 地址和大小与设计规范一致\n'
        '# 示例: "Region 0: Memory at fc000000 (64-bit) [size=64K]"\n'
        '\n'
        '# 5. 如果是网卡驱动，确认网络接口\n'
        'ip link show\n'
        '# 预期: 能看到 DPU 创建的网络接口（如 eth0）'
    )
    add_checklist(doc, [
        'lspci 显示正确的 VID/DID/Class',
        '驱动 probe 成功（dmesg 无报错）',
        '驱动创建的设备节点/网络接口存在',
        'BAR 空间映射正确（地址和大小与规范一致）',
        '中断向量分配成功（cat /proc/interrupts）',
    ])

    doc.add_paragraph('Probe 失败排查流程：')
    add_code_block(doc,
        '# 如果 lspci 看不到设备:\n'
        '#   -> 回到方案一阶段 1, 检查 SHM 通信\n'
        '\n'
        '# 如果 lspci 能看到但驱动未绑定:\n'
        '#   -> 检查 VID/DID 是否三处一致\n'
        'lspci -n  # 显示数字 ID\n'
        '# 对比驱动中的 PCI_DEVICE(vid, did)\n'
        '\n'
        '# 如果驱动 probe 报错:\n'
        '#   -> 查看详细错误\n'
        'dmesg | grep -i "error\\|fail\\|unable"\n'
        '#   -> 常见: BAR 大小不匹配, CONFIG 缺少依赖'
    )

    # 阶段 4: 驱动功能
    doc.add_heading('阶段 4: 驱动功能验证（新增）', level=3)
    doc.add_paragraph('目标：通过驱动提供的标准接口验证各项功能。')

    doc.add_paragraph('4a. 网卡驱动验证：')
    add_code_block(doc,
        '# 链路状态\n'
        'ethtool eth0\n'
        '# 预期: Link detected: yes, Speed: 10000Mb/s (或设计速率)\n'
        '\n'
        '# 驱动信息\n'
        'ethtool -i eth0\n'
        '# 预期: driver: dpu_net, version: x.x.x\n'
        '\n'
        '# 统计信息\n'
        'ethtool -S eth0\n'
        '# 预期: tx_packets/rx_packets 等计数器可读\n'
        '\n'
        '# 配置网络\n'
        'ip addr add 10.0.0.2/24 dev eth0\n'
        'ip link set eth0 up\n'
        '\n'
        '# 查看接口状态\n'
        'ip addr show eth0\n'
        '# 预期: state UP, 地址 10.0.0.2/24'
    )

    doc.add_paragraph('4b. 块设备驱动验证：')
    add_code_block(doc,
        '# 查看块设备\n'
        'lsblk\n'
        '# 预期: 能看到 DPU 块设备\n'
        '\n'
        '# 读测试\n'
        'dd if=/dev/dpu_blk0 of=/dev/null bs=4k count=1\n'
        '\n'
        '# 写测试\n'
        'dd if=/dev/zero of=/dev/dpu_blk0 bs=4k count=1'
    )

    doc.add_paragraph('4c. 自定义驱动验证：')
    add_code_block(doc,
        '# 使用驱动附带的测试工具\n'
        './dpu_test_tool --selftest\n'
        './dpu_test_tool --reg-dump\n'
        './dpu_test_tool --dma-test --size=4096'
    )

    add_checklist(doc, [
        '驱动正确报告设备信息（ethtool -i）',
        '链路状态正确（ethtool 或 carrier 文件）',
        'DMA 映射成功（dmesg 无 IOMMU/SWIOTLB 错误）',
        '中断处理正常（/proc/interrupts 计数增长）',
        '基本数据收发功能正常',
    ])

    # 阶段 5: DMA
    doc.add_heading('阶段 5: DMA 通路验证（通过驱动）', level=3)
    doc.add_paragraph('目标：验证驱动的 DMA 映射与硬件 DMA 引擎协同工作。')
    add_code_block(doc,
        '# DMA 数据流（通过驱动）:\n'
        '# 应用层 -> socket/ioctl -> 驱动 -> dma_map_single()\n'
        '#   -> PCI DMA -> QEMU cosim_pcie_rc -> SHM\n'
        '#   -> VCS bridge -> DPU DMA Engine\n'
        '#   -> DPU 处理 -> DMA Write 回 Guest 内存\n'
        '#   -> MSI 中断 -> 驱动 ISR -> 完成通知\n'
        '\n'
        '# 网卡驱动 DMA 验证:\n'
        '# 发送大包（触发 scatter-gather DMA）\n'
        'ip link set eth0 mtu 9000  # 如果支持 Jumbo Frame\n'
        'dd if=/dev/urandom bs=8000 count=1 | nc -u 10.0.0.1 9999\n'
        '\n'
        '# 检查 DMA 错误:\n'
        'dmesg | grep -i "dma\\|iommu\\|swiotlb"\n'
        '# 预期: 无错误信息'
    )

    add_note(doc,
        '如果 Guest 内核启用了 IOMMU（CONFIG_IOMMU_SUPPORT=y），DMA 地址可能经过 '
        'IOMMU 重映射。CoSim 的 SHM DMA 通路使用物理地址，需确保 IOMMU 配置正确'
        '或在 Guest cmdline 加 "iommu=off"。'
    )

    # 阶段 6: 端到端
    doc.add_heading('阶段 6: 驱动级端到端验证', level=3)
    doc.add_paragraph(
        '目标：通过驱动的标准接口进行完整的应用层端到端测试。'
        '这是方案二的最高验证级别，替代方案一中使用裸工具的阶段 4。'
    )

    doc.add_paragraph('6a. 基本连通性：')
    add_code_block(doc,
        '# 终端 1 (Host 侧): 配置 TAP bridge 或 ETH receiver\n'
        './eth_tap_bridge --eth-shm /cosim_eth_a --tap cosim0 &\n'
        'ip addr add 10.0.0.1/24 dev cosim0\n'
        'ip link set cosim0 up\n'
        '\n'
        '# 终端 2 (Guest 内): 通过驱动发包\n'
        'ip link set eth0 up\n'
        'ip addr add 10.0.0.2/24 dev eth0\n'
        '\n'
        '# ICMP 测试\n'
        'ping -c 5 10.0.0.1\n'
        '# 预期: 5 packets transmitted, 5 received'
    )

    doc.add_paragraph('6b. TCP/UDP 吞吐测试：')
    add_code_block(doc,
        '# Host 侧:\n'
        'iperf3 -s\n'
        '\n'
        '# Guest 内:\n'
        '# TCP 测试\n'
        'iperf3 -c 10.0.0.1 -t 10\n'
        '\n'
        '# UDP 测试\n'
        'iperf3 -c 10.0.0.1 -u -b 1G -t 10\n'
        '\n'
        '# 预期: 能建立连接并传输数据\n'
        '# 注意: CoSim 环境下吞吐量远低于真实硬件，关注正确性而非性能'
    )

    doc.add_paragraph('6c. 驱动专用测试：')
    add_code_block(doc,
        '# 使用驱动附带的测试套件\n'
        './dpu_perf_test --tx-pkts 1000 --pkt-size 64\n'
        './dpu_perf_test --tx-pkts 100 --pkt-size 9000  # Jumbo\n'
        '\n'
        '# 长时间稳定性测试\n'
        './dpu_stress_test --duration 300 --threads 4\n'
        '\n'
        '# 检查驱动统计\n'
        'ethtool -S eth0 | grep -E "tx_|rx_|err"'
    )

    add_checklist(doc, [
        'ping 连通（ICMP 端到端）',
        'TCP 连接建立且数据传输正确',
        'UDP 数据无丢包或在可接受范围',
        '驱动统计计数器与实际收发一致',
        '无 dmesg 错误/警告',
        '长时间运行无内存泄漏或 crash',
    ])

    doc.add_heading('4.6 驱动调试技巧', level=2)
    add_table(doc,
        ['场景', '调试方法', '命令/操作'],
        [
            ['驱动 probe 失败',
             '查看详细内核日志',
             'dmesg | grep -i "dpu|error|fail"'],
            ['VID/DID 不匹配',
             '对比三处 ID',
             'lspci -n 显示设备 ID\ngrep PCI_DEVICE driver/*.c'],
            ['MMIO 读写无响应',
             'VCS 波形 + adapter $display',
             '在 dpu_adapter.sv 添加:\n$display("TLP: type=%0d addr=%0h", ...)'],
            ['DMA 传输错误',
             '驱动 pr_debug + VCS 波形',
             'echo 8 > /proc/sys/kernel/printk\ndynamic_debug 开启驱动 debug'],
            ['中断未触发',
             '检查 MSI 配置',
             'cat /proc/interrupts | grep dpu\nVCS 侧确认 isr_set 信号'],
            ['驱动 panic',
             '开启最大日志级别',
             'Guest cmdline 加: loglevel=7\n或: echo 8 > /proc/sys/kernel/printk'],
            ['version magic 错误',
             '内核版本不匹配',
             '重新用正确 KDIR 编译驱动'],
            ['DMA 地址越界',
             '检查 DMA 位宽',
             '驱动: dma_set_mask(dev, DMA_BIT_MASK(64))\n确保 CoSim 支持 64-bit'],
        ]
    )

    doc.add_heading('4.7 常见问题与排查', level=2)
    add_table(doc,
        ['现象', '原因', '解决方案'],
        [
            ['insmod 报 version magic 错误',
             '驱动编译用的内核头文件与 Guest vmlinuz 版本不匹配',
             '使用与 vmlinuz 完全匹配的 KDIR 重新编译驱动'],
            ['insmod 报 Unknown symbol',
             '依赖的内核模块未加载',
             'modinfo 查看 depends，先 insmod 依赖模块'],
            ['lspci 能看到设备但驱动不 probe',
             'VID/DID 不匹配',
             '检查三处 ID 是否一致'],
            ['驱动 probe 但网口不出现',
             '驱动 register_netdev() 失败',
             'dmesg 查看错误，可能是 BAR 大小不匹配'],
            ['DMA 映射失败',
             'IOMMU 或 SWIOTLB 问题',
             'Guest cmdline 加 "iommu=off swiotlb=force"'],
            ['中断风暴（系统卡死）',
             'ISR 未正确清除中断源',
             '检查驱动 ISR 是否写了正确的中断清除寄存器'],
            ['ping 不通',
             'ARP 未正确响应',
             'tcpdump -i eth0 检查 ARP 包，确认 RX 路径'],
            ['ethtool 报 Operation not supported',
             '驱动未实现 ethtool_ops',
             '非阻塞性问题，可忽略'],
        ]
    )

    doc.add_page_break()

    # ==================== 5. 两方案对比与选型建议 ====================
    doc.add_heading('5. 两方案对比与选型建议', level=1)

    doc.add_heading('5.1 对比总结', level=2)
    add_table(doc,
        ['维度', '方案一：纯 RTL', '方案二：RTL + 驱动'],
        [
            ['部署复杂度', '** (低)', '**** (中)'],
            ['测试深度', '寄存器级 + DMA 级', '应用层端到端'],
            ['调试可见性', '高（直接看波形）', '中（波形 + 驱动日志）'],
            ['验证覆盖率', '中（硬件功能验证）', '高（软硬件协同验证）'],
            ['前置条件', 'RTL 可用即可', 'RTL + 驱动 + 内核源码'],
            ['适用阶段', 'RTL 开发早期 / 驱动未就绪', '驱动已具备基本功能'],
            ['改动范围', '3 个文件', '6+ 个文件'],
            ['问题定位效率', '高（问题一定在 RTL）', '中（需区分 RTL 还是驱动问题）'],
        ]
    )

    doc.add_heading('5.2 选型建议', level=2)
    doc.add_paragraph('推荐的渐进式迁移路径：')
    doc.add_paragraph(
        '第一阶段：先用方案一完成纯 RTL 集成，验证 DPU 的 PCIe EP 基本功能'
        '（Config Space、BAR、DMA、中断），确保硬件接口无误。',
        style='List Bullet'
    )
    doc.add_paragraph(
        '第二阶段：在方案一所有测试 PASS 后，切换到方案二，加载驱动进行端到端验证。'
        '此时如果出现问题，可以确定是驱动层面的问题（因为硬件层面已在方案一中验证通过）。',
        style='List Bullet'
    )
    doc.add_paragraph(
        '第三阶段：基于方案二的环境，运行驱动的完整测试套件，进行功能回归和稳定性测试。',
        style='List Bullet'
    )

    add_note(doc,
        '不建议跳过方案一直接执行方案二。如果驱动 probe 失败，'
        '在方案一未验证的情况下很难判断是 RTL 问题还是驱动问题。'
    )

    doc.add_page_break()

    # ==================== 6. 文件修改清单总览 ====================
    doc.add_heading('6. 文件修改清单总览', level=1)

    doc.add_heading('6.1 方案一文件修改', level=2)
    add_table(doc,
        ['文件', '操作', '说明'],
        [
            ['vcs-tb/pcie_ep_stub.sv', '保留不编译', '作为接口参考'],
            ['vcs-tb/dpu_adapter.sv', '新建', '接口适配层（核心工作）'],
            ['vcs-tb/tb_top.sv', '修改', '替换 EP 实例为 dpu_adapter'],
            ['build_dpu_simv.sh', '新建', 'DPU 版 VCS 编译脚本'],
            ['setup.sh', '无修改', 'QEMU/Bridge 构建不变'],
            ['qemu-plugin/cosim_pcie_rc.c', '无修改', 'SHM 协议不变'],
            ['bridge/vcs/bridge_vcs.c', '无修改', 'DPI-C 接口不变'],
        ]
    )

    doc.add_heading('6.2 方案二额外文件修改', level=2)
    doc.add_paragraph('在方案一的基础上，还需修改：')
    add_table(doc,
        ['文件', '操作', '说明'],
        [
            ['qemu-plugin/cosim_pcie_rc.c', '修改', 'VID/DID/BAR 大小匹配 DPU 驱动'],
            ['scripts/guest_init_dpu.sh', '新建', 'DPU 驱动加载 + 接口配置脚本'],
            ['initramfs', '重新打包', '包含 DPU 驱动 .ko + 固件 + init 修改'],
            ['rebuild_initramfs_with_driver.sh', '新建', 'Initramfs 打包自动化脚本'],
            ['Guest 内核 .config', '可能修改', '确保驱动依赖的 CONFIG 开启'],
        ]
    )

    doc.add_page_break()

    # ==================== 7. 附录 ====================
    doc.add_heading('7. 附录', level=1)

    doc.add_heading('7.1 完整接口信号列表', level=2)
    doc.add_paragraph('以下是 pcie_ep_stub.sv 模块的完整端口定义，dpu_adapter.sv 必须实现相同接口：')
    add_code_block(doc,
        'module dpu_adapter (\n'
        '    // 时钟与复位\n'
        '    input  logic        clk,            // 系统时钟\n'
        '    input  logic        rst_n,          // 异步复位（低有效）\n'
        '\n'
        '    // TLP 请求 (Host -> EP)\n'
        '    input  logic        tlp_valid,      // 请求有效\n'
        '    input  logic [2:0]  tlp_type,       // 0:MRd 1:MWr 2:CfgRd 3:CfgWr 4:DMA_CPL\n'
        '    input  logic [63:0] tlp_addr,       // 目标地址\n'
        '    input  logic [31:0] tlp_wdata,      // 写数据\n'
        '    input  logic [15:0] tlp_len,        // 长度（字节）\n'
        '    input  logic [7:0]  tlp_tag,        // 事务标识\n'
        '\n'
        '    // TLP 完成 (EP -> Host)\n'
        '    output logic        cpl_valid,      // 完成有效\n'
        '    output logic [7:0]  cpl_tag,        // 对应请求 tag\n'
        '    output logic [31:0] cpl_rdata,      // 读返回数据\n'
        '    output logic        cpl_status,     // 0:成功 1:失败\n'
        '\n'
        '    // Virtio 通知 (EP -> Host, 可选)\n'
        '    output logic        notify_valid,   // 通知有效\n'
        '    output logic [15:0] notify_queue,   // 队列编号\n'
        '\n'
        '    // 中断 (EP -> Host)\n'
        '    output logic        isr_set         // 中断请求\n'
        ');'
    )

    doc.add_heading('7.2 DPI-C 函数参考', level=2)
    doc.add_paragraph('bridge_vcs.sv 包中声明的 DPI-C 函数，供 tb_top.sv 调用：')
    add_table(doc,
        ['函数', '方向', '说明'],
        [
            ['bridge_vcs_init(name)', 'C->SV', '初始化 SHM，参数为共享内存名'],
            ['bridge_vcs_poll_tlp(...)', 'C->SV', '从 SHM 请求队列取 TLP，返回 0=有请求'],
            ['bridge_vcs_send_completion(...)', 'SV->C', '将完成数据写入 SHM 完成队列'],
            ['bridge_vcs_dma_read_sync(...)', 'SV->C', '同步 DMA 读：从 Guest 内存读数据'],
            ['bridge_vcs_dma_write_sync(...)', 'SV->C', '同步 DMA 写：向 Guest 内存写数据'],
            ['bridge_vcs_raise_msi(vector)', 'SV->C', '触发 MSI 中断'],
            ['vcs_eth_mac_init_dpi(name)', 'C->SV', '初始化以太网 SHM'],
            ['vcs_eth_mac_send_frame_dpi(...)', 'SV->C', '发送以太网帧到 SHM'],
            ['vcs_eth_mac_poll_frame_dpi(...)', 'C->SV', '从 SHM 接收以太网帧'],
            ['vcs_vq_configure(...)', 'SV->C', '配置 Virtqueue 描述符/可用/已用环地址'],
            ['vcs_vq_process_tx(...)', 'SV->C', '处理 TX Virtqueue（DMA 读描述符+数据）'],
            ['vcs_vq_process_rx(...)', 'SV->C', '处理 RX Virtqueue（DMA 写数据到描述符）'],
        ]
    )

    doc.add_heading('7.3 SHM 内存布局', level=2)
    doc.add_paragraph('共享内存总大小 4MB，布局如下：')
    add_table(doc,
        ['区域', '偏移', '大小', '说明'],
        [
            ['ctrl', '0x00000000', '4KB', '控制区：magic/version/ready 标志/sim_time'],
            ['req_ring', '0x00001000', '256KB', 'TLP 请求环形缓冲区（Host->EP）'],
            ['cpl_ring', '0x00041000', '256KB', 'TLP 完成环形缓冲区（EP->Host）'],
            ['dma_req_ring', '0x00081000', '64KB', 'DMA 请求队列'],
            ['dma_cpl_ring', '0x00091000', '64KB', 'DMA 完成队列'],
            ['msi_ring', '0x000A1000', '4KB', 'MSI 中断事件队列'],
            ['dma_buf', '0x000A2000', '~3.4MB', 'DMA 数据缓冲区'],
        ]
    )

    # ==================== 尾页 ====================
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end.add_run('-- 文档结束 --')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return doc


if __name__ == '__main__':
    import os
    doc = create_document()
    out_dir = os.path.join(os.path.dirname(__file__), 'cosim-platform')
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, 'CoSim_DPU_Migration_Integration.docx')
    doc.save(out_path)
    print(f'Document saved: {out_path}')
    print(f'Size: {os.path.getsize(out_path) / 1024:.1f} KB')
