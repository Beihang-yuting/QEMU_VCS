#!/usr/bin/env python3
"""
生成 CoSim Platform 完整部署与测试流程 Word 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime


def add_code_block(doc, code):
    """添加代码块（灰色背景）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F0F0',
        qn('w:val'): 'clear'
    })
    pPr.append(shd)


def add_table_row(table, cells_text):
    """添加表格行"""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        row.cells[i].text = text
    return row


def make_header_bold(table):
    """设置表头加粗"""
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True


def create_document():
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ===== 封面 =====
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('QEMU-VCS CoSim Platform\n完整部署与测试流程')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xc4)

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'版本: 1.0\n日期: {datetime.date.today().strftime("%Y-%m-%d")}')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ===== 目录 =====
    doc.add_heading('目录', level=1)
    toc_items = [
        '1. 项目概述',
        '2. 系统架构',
        '3. 环境要求',
        '4. 部署流程',
        '  4.1 获取源码',
        '  4.2 配置环境变量 (config.env)',
        '  4.3 运行安装脚本 (setup.sh)',
        '  4.4 安装步骤详解',
        '5. 测试流程',
        '  5.1 测试阶段总览',
        '  5.2 Phase 1-3: 单 VCS 测试',
        '  5.3 Phase 4: 双 VCS 环回测试',
        '  5.4 Phase 5: 双 VCS 端到端网络测试',
        '  5.5 TAP Bridge 模式测试',
        '6. 统一命令行工具 (cosim.sh)',
        '7. 故障排查',
        '8. 附录',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # ===== 1. 项目概述 =====
    doc.add_heading('1. 项目概述', level=1)
    doc.add_paragraph(
        'QEMU-VCS CoSim Platform 是一个软硬件协同仿真平台，通过共享内存（SHM）实现 '
        'QEMU 虚拟机与 Synopsys VCS 仿真器之间的高速 PCIe TLP 通信。平台支持完整的 '
        'virtio-net 网络设备仿真，包括 DMA 传输、MSI 中断、虚拟队列管理等功能，'
        '实现了 Guest 操作系统的 TCP/UDP 网络通信。'
    )

    doc.add_heading('核心功能', level=2)
    features = [
        'PCIe TLP 级别的软硬件协同仿真（QEMU <-> VCS 共享内存）',
        'VirtIO-Net 设备完整实现（Virtqueue + DMA + MSI）',
        '以太网 MAC 层仿真与 SHM 转发',
        '双 VCS 模式：两个 Guest 直接网络通信',
        'TAP Bridge 模式：Guest 与宿主机网络通信',
        '一键部署脚本，支持多种 Linux 发行版',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    # ===== 2. 系统架构 =====
    doc.add_heading('2. 系统架构', level=1)

    doc.add_heading('2.1 整体架构', level=2)
    doc.add_paragraph(
        '系统由以下核心组件组成：\n\n'
        '1) QEMU (cosim-pcie-rc 设备): 提供 PCIe Root Complex 模拟，通过 SHM 与 VCS 通信\n'
        '2) VCS (bridge_vcs): SystemVerilog testbench 模拟 PCIe EP 设备，处理 TLP 包\n'
        '3) Bridge 库: C 库实现 SHM 管理、Ring Buffer、DMA、中断等核心功能\n'
        '4) ETH SHM: 以太网共享内存层，用于双 VCS 间或 VCS 与 TAP bridge 间帧转发\n'
        '5) eth_tap_bridge: 用户态桥接程序，连接 ETH SHM 与 Linux TAP 设备'
    )

    doc.add_heading('2.2 数据通路', level=2)

    doc.add_paragraph('Phase 4/5（双 VCS 模式）：')
    add_code_block(doc,
        'Guest1 (eth0) -> QEMU1 (cosim-pcie-rc) -> SHM1 -> VCS-A (ETH SHM) -> VCS-B -> SHM2 -> QEMU2 -> Guest2 (eth0)')

    doc.add_paragraph('TAP Bridge 模式：')
    add_code_block(doc,
        'Guest (eth0) -> QEMU (cosim-pcie-rc) -> SHM -> VCS (ETH SHM) -> eth_tap_bridge -> TAP (cosim0) -> Host')

    # ===== 3. 环境要求 =====
    doc.add_heading('3. 环境要求', level=1)

    doc.add_heading('3.1 硬件要求', level=2)
    for item in ['x86_64 处理器', '内存 >= 4GB（推荐 8GB+）', '磁盘空间 >= 10GB（含 QEMU 编译）']:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.2 软件要求', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = '组件'
    table.rows[0].cells[1].text = '最低版本'
    table.rows[0].cells[2].text = '说明'
    make_header_bold(table)

    for row_data in [
        ('Linux', 'CentOS 7+ / Ubuntu 18.04+', '支持 POSIX SHM'),
        ('GCC', '4.8+', 'C11 原子操作兼容'),
        ('CMake', '>= 3.16', '构建系统'),
        ('Python3', '>= 3.6', 'QEMU 构建依赖'),
        ('Meson + Ninja', '-', 'QEMU 构建依赖'),
        ('VCS', 'Q-2020.03+', 'Synopsys 仿真器（可选）'),
        ('QEMU 源码', '9.2.0', '自动下载或手动放置'),
    ]:
        add_table_row(table, row_data)

    # ===== 4. 部署流程 =====
    doc.add_heading('4. 部署流程', level=1)

    doc.add_heading('4.1 获取源码', level=2)
    doc.add_paragraph('将 cosim-platform 目录拷贝到目标机器：')
    add_code_block(doc, 'scp -r cosim-platform/ user@target:~/workspace/')

    doc.add_heading('4.2 配置环境变量 (config.env)', level=2)
    doc.add_paragraph('编辑 config.env 文件，适配目标机器环境。主要配置项：')

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '变量名'
    table.rows[0].cells[1].text = '默认值'
    table.rows[0].cells[2].text = '说明'
    make_header_bold(table)

    for row_data in [
        ('VCS_HOME', '(空)', 'VCS 安装目录，留空自动搜索'),
        ('SNPSLMD_LICENSE_FILE', '/opt/synopsys/license/license.dat', 'VCS 许可证路径'),
        ('QEMU_VERSION', 'v9.2.0', 'QEMU 版本号'),
        ('GUEST1_IP / GUEST2_IP', '10.0.0.1 / 10.0.0.2', 'Guest 网络 IP'),
        ('GUEST_MEMORY', '256M', 'Guest 内存大小'),
        ('PHASE5_TIMEOUT', '300', 'Phase 5 超时（秒）'),
    ]:
        add_table_row(table, row_data)

    doc.add_heading('4.3 运行安装脚本 (setup.sh)', level=2)
    doc.add_paragraph('一键安装命令：')
    add_code_block(doc, 'cd ~/workspace/cosim-platform\nbash setup.sh')
    doc.add_paragraph('如果 QEMU 源码已存在于其他位置，可通过环境变量指定：')
    add_code_block(doc, 'QEMU_SRC_DIR=~/workspace/qemu-9.2.0 bash setup.sh')

    doc.add_heading('4.4 安装步骤详解', level=2)

    steps = [
        ('[1/9] 加载配置文件', '读取 config.env，设置 VCS 许可证、网络参数等环境变量'),
        ('[2/9] 检测 OS 并安装依赖', '自动检测 Debian/CentOS，安装 gcc, cmake, meson 等。无 sudo 时跳过并提示'),
        ('[3/9] 编译 Bridge 库', 'CMake 编译 libcosim_bridge.so（SHM, Ring Buffer, DMA, 中断等核心库）'),
        ('[4/9] 编译 QEMU', '注入 cosim_pcie_rc 设备到 QEMU 源码树，编译 qemu-system-x86_64'),
        ('[5/9] 编译 VCS simv', '使用 VCS 编译器生成仿真可执行文件 simv（含 DPI-C bridge 代码）'),
        ('[6/9] 编译 eth_tap_bridge', '编译用户态 TAP 桥接工具'),
        ('[7/9] 构建 initramfs', '构建 Guest 启动镜像（含 virtio_net 驱动、busybox、测试脚本）'),
        ('[8/9] 运行单元测试', '执行 Ring Buffer、SHM Layout、DMA Manager 等组件的单元测试'),
        ('[9/9] 安装摘要', '列出所有构建产物状态和警告汇总'),
    ]
    for step_name, step_desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(step_name + ': ')
        run.font.bold = True
        p.add_run(step_desc)

    doc.add_paragraph()
    doc.add_paragraph('安装完成后的产物：')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '产物'
    table.rows[0].cells[1].text = '路径'
    table.rows[0].cells[2].text = '说明'
    make_header_bold(table)

    for row_data in [
        ('libcosim_bridge.so', 'build/bridge/', 'SHM 通信核心库'),
        ('qemu-system-x86_64', 'third_party/qemu/build/', '含 cosim-pcie-rc 设备的 QEMU'),
        ('simv', 'vcs-tb/sim_build/', 'VCS 仿真可执行文件'),
        ('eth_tap_bridge', 'tools/', 'TAP 桥接工具'),
        ('initramfs-*.cpio.gz', 'images/', 'Guest 启动镜像'),
    ]:
        add_table_row(table, row_data)

    doc.add_page_break()

    # ===== 5. 测试流程 =====
    doc.add_heading('5. 测试流程', level=1)

    doc.add_heading('5.1 测试阶段总览', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '阶段'
    table.rows[0].cells[1].text = '描述'
    table.rows[0].cells[2].text = '验证内容'
    table.rows[0].cells[3].text = '命令'
    make_header_bold(table)

    for row_data in [
        ('Phase 1', 'Config Space', 'PCIe 配置空间读写', './cosim.sh test phase1'),
        ('Phase 2', 'DMA 读写', 'Host-Guest DMA 数据传输', './cosim.sh test phase2'),
        ('Phase 3', 'MSI 中断', 'MSI 中断触发与接收', './cosim.sh test phase3'),
        ('Phase 4', '环回测试', 'VirtIO-Net + DMA + 双 VCS 环回', './cosim.sh test phase4'),
        ('Phase 5', '端到端网络', 'Guest 间 ping/nc/iperf3', './cosim.sh test phase5'),
        ('TAP', '宿主机桥接', 'Guest 与 Host 网络通信', './cosim.sh test tap'),
    ]:
        add_table_row(table, row_data)

    doc.add_heading('5.2 Phase 1-3: 单 VCS 测试', level=2)
    doc.add_paragraph(
        '这三个阶段使用单个 QEMU + 单个 VCS 实例，验证基本 PCIe 通信功能：'
    )
    doc.add_paragraph('Phase 1 验证 PCIe 配置空间的 BAR 映射和寄存器读写', style='List Bullet')
    doc.add_paragraph('Phase 2 验证 DMA 读写通路，QEMU 发送 MRd/MWr TLP，VCS 返回 Completion', style='List Bullet')
    doc.add_paragraph('Phase 3 验证 MSI 中断从 VCS EP 到 QEMU Guest 的完整通路', style='List Bullet')
    add_code_block(doc, './cosim.sh test phase1\n./cosim.sh test phase2\n./cosim.sh test phase3')

    doc.add_heading('5.3 Phase 4: 双 VCS 环回测试', level=2)
    doc.add_paragraph(
        '使用两个 QEMU + 两个 VCS 实例，通过 ETH SHM 连接，验证 VirtIO-Net 设备的完整功能。'
        '两个 Guest 各配置一个 eth0 接口，通过 virtio-net 驱动发送以太网帧，'
        '经 DMA -> VCS -> ETH SHM -> VCS -> DMA 环回路径传输。'
    )
    add_code_block(doc, './cosim.sh test phase4')
    doc.add_paragraph('预期结果：')
    doc.add_paragraph('Guest1 (10.0.0.1) <-> Guest2 (10.0.0.2) 互 ping 成功', style='List Bullet')
    doc.add_paragraph('RX/TX 包计数正确，无丢包', style='List Bullet')

    doc.add_heading('5.4 Phase 5: 双 VCS 端到端网络测试', level=2)
    doc.add_paragraph('在 Phase 4 基础上增加 TCP/UDP 应用层测试：')
    doc.add_paragraph('nc (netcat) 文本传输测试', style='List Bullet')
    doc.add_paragraph('iperf3 TCP 吞吐量测试', style='List Bullet')
    doc.add_paragraph('iperf3 UDP 吞吐量测试', style='List Bullet')
    add_code_block(doc, './cosim.sh test phase5')

    doc.add_paragraph('Phase 5 测试结果（实测）：')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '测试项'
    table.rows[0].cells[1].text = '结果'
    table.rows[0].cells[2].text = '数据'
    make_header_bold(table)
    for row_data in [
        ('Ping 20 packets', 'PASS', '20/20, 0% loss'),
        ('nc 文本传输', 'PASS', '双向传输成功'),
        ('iperf3 TCP', 'PASS', '9.74 Mbps, 0 retransmissions'),
        ('iperf3 UDP', 'PASS', '8.98 Mbps'),
    ]:
        add_table_row(table, row_data)

    doc.add_heading('5.5 TAP Bridge 模式测试', level=2)
    doc.add_paragraph(
        '单 QEMU + 单 VCS + eth_tap_bridge 模式，Guest 通过虚拟网卡与宿主机 TAP 设备通信。'
        '适用于需要 Guest 访问外部网络的场景。'
    )
    add_code_block(doc,
        '# 架构：\n'
        '# Guest eth0 (10.0.0.2) <-> QEMU <-> SHM <-> VCS <-> ETH SHM <-> eth_tap_bridge <-> TAP cosim0 (10.0.0.1)\n\n'
        './cosim.sh test tap')
    doc.add_paragraph('预期结果：')
    doc.add_paragraph('Guest (10.0.0.2) -> Host TAP (10.0.0.1) ping PASS', style='List Bullet')
    doc.add_paragraph('Host TAP (10.0.0.1) -> Guest (10.0.0.2) ping PASS', style='List Bullet')

    doc.add_page_break()

    # ===== 6. 统一命令行工具 =====
    doc.add_heading('6. 统一命令行工具 (cosim.sh)', level=1)
    doc.add_paragraph('cosim.sh 是平台的统一入口脚本，支持测试、启动、状态查看、日志查看等功能。')

    doc.add_heading('6.1 命令总览', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '命令'
    table.rows[0].cells[1].text = '用法'
    table.rows[0].cells[2].text = '说明'
    make_header_bold(table)
    for row_data in [
        ('test', './cosim.sh test <phase>', '运行指定测试阶段'),
        ('start', './cosim.sh start <component>', '启动单个组件'),
        ('status', './cosim.sh status', '查看运行状态'),
        ('log', './cosim.sh log <component>', '查看组件日志'),
        ('clean', './cosim.sh clean', '清理 SHM 和临时文件'),
        ('info', './cosim.sh info', '显示系统信息'),
        ('help', './cosim.sh help', '显示帮助信息'),
    ]:
        add_table_row(table, row_data)

    doc.add_heading('6.2 常用示例', level=2)
    add_code_block(doc,
        '# 运行全部测试\n'
        './cosim.sh test all\n\n'
        '# 仅运行 Phase 5\n'
        './cosim.sh test phase5\n\n'
        '# 运行 TAP 桥接测试\n'
        './cosim.sh test tap\n\n'
        '# 查看状态\n'
        './cosim.sh status\n\n'
        '# 清理资源\n'
        './cosim.sh clean')

    # ===== 7. 故障排查 =====
    doc.add_heading('7. 故障排查', level=1)

    issues = [
        ('VCS 许可证错误',
         '检查 SNPSLMD_LICENSE_FILE 环境变量指向正确的 license.dat 文件。\n'
         '确认许可证服务器正常运行: lmstat -a -c /opt/synopsys/license/license.dat'),
        ('QEMU 编译失败 - 缺少依赖',
         '安装 glib2-devel, pixman-devel, meson, ninja-build。\n'
         'CentOS: yum install glib2-devel pixman-devel\n'
         'Ubuntu: apt install libglib2.0-dev libpixman-1-dev'),
        ('SHM 创建失败',
         '检查 /dev/shm 目录权限和可用空间。\n'
         '清理旧 SHM: rm -f /dev/shm/cosim*'),
        ('Guest 网络不通',
         '确认 Guest init 脚本中的 IP 和 MAC 地址与 config.env 一致。\n'
         '检查 VCS 日志中的 ETH 转发记录。\n'
         '确认静态 ARP 表项正确配置。'),
        ('TAP 设备创建失败',
         '检查 /dev/net/tun 设备权限（需要 rw 权限）。\n'
         '确认 ip 命令可用: /sbin/ip 或 /usr/sbin/ip'),
        ('BQL 死锁（iperf3 卡住）',
         '确保使用了 BQL 非阻塞 MSI 修复版本的 QEMU 插件。\n'
         'cosim_pcie_rc.c 中 MSI 应通过 BH (Bottom Half) 机制发送，避免 irq_poller 线程持锁。'),
        ('GCC 4.8 编译错误',
         'setup.sh 已自动处理 stdatomic.h 兼容性（通过 compat_atomic.h）。\n'
         'VCS 编译自动添加 -std=gnu99 和 -timescale=1ns/1ps 参数。'),
    ]
    for title, desc in issues:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
        doc.add_paragraph(desc)

    # ===== 8. 附录 =====
    doc.add_heading('8. 附录', level=1)

    doc.add_heading('8.1 目录结构', level=2)
    add_code_block(doc,
        'cosim-platform/\n'
        '  config.env           # 环境配置文件\n'
        '  setup.sh             # 一键安装脚本\n'
        '  cosim.sh             # 统一 CLI 工具\n'
        '  build_initramfs.sh   # initramfs 构建脚本\n'
        '  CMakeLists.txt       # CMake 构建配置\n'
        '  bridge/              # 核心 Bridge 库\n'
        '    common/            #   SHM, Ring Buffer, DMA, 原子操作兼容\n'
        '    qemu/              #   QEMU 侧桥接代码\n'
        '    vcs/               #   VCS 侧桥接代码 (DPI-C)\n'
        '    eth/               #   以太网 MAC 层\n'
        '  qemu-plugin/         # QEMU cosim-pcie-rc 设备插件\n'
        '  vcs-tb/              # VCS testbench (SystemVerilog)\n'
        '  tools/               # eth_tap_bridge 等辅助工具\n'
        '  scripts/             # 测试脚本和 Guest init 脚本\n'
        '  tests/               # 单元测试和集成测试\n'
        '  images/              # 构建产物（内核、initramfs）')

    doc.add_heading('8.2 关键技术参数', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '参数'
    table.rows[0].cells[1].text = '值'
    make_header_bold(table)
    for row_data in [
        ('PCIe SHM 大小', '4 MB'),
        ('TLP Ring Buffer 容量', '512 entries'),
        ('ETH SHM Frame Ring 深度', '256 frames'),
        ('最大以太网帧长度', '1514 bytes'),
        ('Guest 内存', '256 MB'),
        ('DMA 区域大小', '1 MB'),
        ('MSI 队列深度', '256'),
        ('iperf3 TCP 吞吐量', '~9.74 Mbps'),
    ]:
        add_table_row(table, row_data)

    # 保存
    output_path = '/home/ubuntu/ryan/software/cosim-platform/CoSim_Platform_部署测试流程.docx'
    doc.save(output_path)
    print(f'Word 文档已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    create_document()
