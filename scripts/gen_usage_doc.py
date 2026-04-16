#!/usr/bin/env python3
"""生成 QEMU-VCS CoSim Platform 使用说明 Word 文档。

Output: cosim-platform/docs/CoSim-Platform-Usage-Guide.docx
"""
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "CoSim-Platform-Usage-Guide.docx"

COLOR_H1 = RGBColor(0x1F, 0x4E, 0x79)
COLOR_H2 = RGBColor(0x2E, 0x75, 0xB6)
COLOR_H3 = RGBColor(0x5B, 0x9B, 0xD5)
COLOR_CODE_BG = "F2F2F2"
COLOR_NOTE = RGBColor(0xBF, 0x8F, 0x00)


def _set_chinese_font(run, name: str = "Microsoft YaHei") -> None:
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = COLOR_H1
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = COLOR_H2
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_H3
    _set_chinese_font(run)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    run.italic = italic
    _set_chinese_font(run)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    _set_chinese_font(run)


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): COLOR_CODE_BG,
    })
    pPr.append(shd)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rFonts.set(qn("w:eastAsia"), "Consolas")
    rFonts.set(qn("w:cs"), "Consolas")


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("注意：")
    run.bold = True
    run.font.color.rgb = COLOR_NOTE
    run.font.size = Pt(10.5)
    _set_chinese_font(run)
    run2 = p.add_run(text)
    run2.font.size = Pt(10.5)
    _set_chinese_font(run2)


def add_table(doc: Document, headers, rows) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        _set_chinese_font(run)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            _set_chinese_font(run)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # ---------- Title page ----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("QEMU-VCS 协同仿真平台")
    t_run.bold = True
    t_run.font.size = Pt(32)
    t_run.font.color.rgb = COLOR_H1
    _set_chinese_font(t_run)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = subtitle.add_run("使用说明书（Usage Guide）")
    s_run.bold = True
    s_run.font.size = Pt(18)
    s_run.font.color.rgb = COLOR_H2
    _set_chinese_font(s_run)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "版本：P3（双节点 ETH 互打 + 链路模型）",
        "日期：2026-04-16",
        "维护者：Beihang-yuting <2965455908@qq.com>",
        "仓库：https://github.com/Beihang-yuting/QEMU_VCS",
    ]:
        r = meta.add_run(line + "\n")
        r.font.size = Pt(11)
        _set_chinese_font(r)

    doc.add_page_break()

    # ---------- 1. 概述 ----------
    add_heading(doc, "1. 平台概述", 1)
    add_para(doc,
             "QEMU-VCS 协同仿真平台是一个将 QEMU 虚拟化环境与 Synopsys VCS RTL 仿真器"
             "打通的异构联合仿真基础设施，用于在 RTL 阶段即可以运行真实操作系统 / 驱动对"
             "待测 ASIC（例如 DPU、智能网卡、加速器）进行软硬件协同验证。")

    add_heading(doc, "1.1 核心能力", 2)
    add_bullet(doc, "PCIe MMIO 读写事务联合仿真（TLP 级）")
    add_bullet(doc, "DMA 数据面双向传输（设备→Host / Host→设备）")
    add_bullet(doc, "MSI 中断注入（VCS RTL → QEMU Guest）")
    add_bullet(doc, "双模式同步：快速模式（事务级）与精确模式（周期级锁步）")
    add_bullet(doc, "运行时模式切换、事务级追踪日志（CSV/JSON）")
    add_bullet(doc, "POSIX 共享内存 + Unix Domain Socket 低开销 IPC")
    add_bullet(doc, "**双节点以太网互打**（P3）：ETH SHM 帧队列、9KB Jumbo、松耦合时间同步")
    add_bullet(doc, "**链路模型**（P3）：丢包率、突发丢包、固定延迟、速率限、流控窗口")
    add_bullet(doc, "**调试工具链**（P4）：cosim_cli REPL、trace_analyzer、launch_dual 编排")

    add_heading(doc, "1.2 系统架构", 2)
    add_para(doc,
             "平台由两类进程构成，通过 POSIX 共享内存（数据面）与 Unix Domain Socket"
             "（控制面）通信：")
    add_bullet(doc, "QEMU 进程：运行 Guest Linux，加载自定义 PCIe RC 设备 cosim-pcie-rc")
    add_bullet(doc, "VCS 仿真进程：运行 RTL 设计（DPU PCIe EP），通过 DPI-C 调用 Bridge")
    add_bullet(doc, "Bridge 库（libcosim_bridge.so / libcosim_bridge_vcs.so）：承载 TLP 请求队列、"
                   "完成队列、DMA 数据区、MSI 事件队列、时钟同步握手")

    doc.add_page_break()

    # ---------- 2. 目录结构 ----------
    add_heading(doc, "2. 仓库目录结构", 1)
    add_code(doc,
             "cosim-platform/\n"
             "├── CMakeLists.txt              # 顶层 CMake\n"
             "├── Makefile                    # 一键 make / make test\n"
             "├── bridge/                     # Bridge 库源码\n"
             "│   ├── common/                 # 平台共用：SHM、环形缓冲、DMA allocator、trace\n"
             "│   ├── qemu/                   # QEMU 侧 Bridge（libcosim_bridge.so）\n"
             "│   └── vcs/                    # VCS 侧 Bridge（libcosim_bridge_vcs.so + .sv）\n"
             "├── qemu-plugin/                # QEMU 自定义 PCIe RC 设备（需装入 QEMU 源码树）\n"
             "│   ├── cosim_pcie_rc.c\n"
             "│   └── cosim_pcie_rc.h\n"
             "├── vcs-tb/                     # VCS 测试平台\n"
             "│   ├── tb_top.sv               # 顶层 TB\n"
             "│   └── pcie_ep_stub.sv         # PCIe EP 简化模型\n"
             "├── scripts/\n"
             "│   ├── setup_cosim_qemu.sh     # 一键集成设备到 QEMU 源码树\n"
             "│   ├── run_cosim.sh            # 启动 QEMU 进程\n"
             "│   └── gen_usage_doc.py        # 生成本使用说明 Word 文档\n"
             "├── tests/\n"
             "│   ├── unit/                   # 4 个单元测试\n"
             "│   └── integration/            # 5 个集成测试\n"
             "└── docs/                       # 生成的使用说明与规划文档")

    doc.add_page_break()

    # ---------- 3. 环境要求 ----------
    add_heading(doc, "3. 环境要求", 1)
    add_heading(doc, "3.1 操作系统与基础工具", 2)
    add_table(doc,
              headers=["组件", "版本要求", "用途"],
              rows=[
                  ["Linux", "Ubuntu 20.04+ / CentOS 8+ (x86_64)", "运行 QEMU + VCS"],
                  ["GCC", ">= 9.0 (支持 C11 / __atomic)", "编译 Bridge"],
                  ["CMake", ">= 3.16", "构建系统"],
                  ["pthread / librt", "glibc 内置", "Bridge 线程与 SHM"],
                  ["QEMU", "7.0+ (含 KVM 可选)", "运行 Guest OS"],
                  ["Synopsys VCS", "2020.03+（含 DPI-C 支持）", "RTL 仿真"],
              ])

    add_heading(doc, "3.2 可选组件", 2)
    add_bullet(doc, "KVM（/dev/kvm）：可选，启用后 Guest 运行性能接近原生")
    add_bullet(doc, "GDB：可选，用于调试 Guest 内核或驱动")
    add_bullet(doc, "Docker：可选，可在容器内完成构建和单元/集成测试（无需真实 VCS）")

    add_heading(doc, "3.3 资源需求", 2)
    add_bullet(doc, "磁盘：约 500 MB（Bridge + QEMU 构建产物，不含 Guest 镜像）")
    add_bullet(doc, "内存：至少 8 GB（含 QEMU Guest 4 GB + VCS 2 GB）")
    add_bullet(doc, "共享内存：64 MB（可配置），通过 /dev/shm 挂载")

    doc.add_page_break()

    # ---------- 4. 快速开始 ----------
    add_heading(doc, "4. 快速开始", 1)

    add_heading(doc, "4.1 克隆仓库", 2)
    add_code(doc, "git clone https://github.com/Beihang-yuting/QEMU_VCS.git\n"
                  "cd QEMU_VCS")

    add_heading(doc, "4.2 构建 Bridge 与测试", 2)
    add_code(doc,
             "# 构建\n"
             "make bridge\n\n"
             "# 运行全部单元+集成测试\n"
             "make test\n\n"
             "# 预期：9/9 测试通过")
    add_para(doc, "测试项（共 17 个）：")
    add_bullet(doc, "单元 P1/P2/P3：test_ring_buffer, test_shm_layout, test_dma_manager, "
                   "test_trace_log, test_eth_shm, test_link_model")
    add_bullet(doc, "集成 PCIe：test_sock_sync, test_bridge_loopback, test_dma_roundtrip, "
                   "test_msi_roundtrip, test_precise_mode")
    add_bullet(doc, "集成 ETH（P3）：test_eth_loopback, test_link_drop, test_mac_stub_e2e, "
                   "test_time_sync_loose")
    add_bullet(doc, "工具（P4）：test_cli_smoke, test_launch_smoke")

    add_heading(doc, "4.3 集成到 QEMU 源码树", 2)
    add_para(doc, "自定义 PCIe RC 设备需要装入 QEMU 源码树后再编译 QEMU：")
    add_code(doc, "# 假设 QEMU 源码位于 ./third_party/qemu\n"
                  "export QEMU_SRC=./third_party/qemu\n"
                  "./scripts/setup_cosim_qemu.sh $QEMU_SRC\n\n"
                  "# 编译 QEMU（按需增加 --target-list 等）\n"
                  "cd $QEMU_SRC\n"
                  "./configure --target-list=x86_64-softmmu --enable-kvm\n"
                  "make -j$(nproc)")
    add_note(doc, "setup_cosim_qemu.sh 会把 cosim_pcie_rc.{c,h} 复制到 hw/net/ 并更新 meson.build。")

    add_heading(doc, "4.4 编译 VCS 侧", 2)
    add_code(doc,
             "vcs -full64 -sverilog \\\n"
             "    -CFLAGS \"-I bridge/common -I bridge/vcs\" \\\n"
             "    -LDFLAGS \"-L build/bridge -lcosim_bridge_vcs -lrt -lpthread\" \\\n"
             "    bridge/vcs/bridge_vcs.sv vcs-tb/*.sv \\\n"
             "    -o simv")

    add_heading(doc, "4.5 启动联合仿真", 2)
    add_para(doc, "在两个终端分别启动 QEMU 与 VCS（QEMU 必须先起，VCS 作为 client 接入）：")
    add_code(doc,
             "# 终端 A：启动 QEMU\n"
             "export GUEST_KERNEL=/path/to/bzImage\n"
             "export GUEST_ROOTFS=/path/to/rootfs.qcow2   # 可选\n"
             "./scripts/run_cosim.sh\n\n"
             "# 终端 B：启动 VCS\n"
             "./simv +SHM_NAME=/cosim0 +SOCK_PATH=/tmp/cosim.sock")
    add_note(doc, "默认 SHM_NAME=/cosim0, SOCK_PATH=/tmp/cosim.sock。可通过环境变量覆盖。")

    doc.add_page_break()

    # ---------- 5. 运行模式 ----------
    add_heading(doc, "5. 运行模式", 1)
    add_table(doc,
              headers=["模式", "同步粒度", "性能", "适用场景"],
              rows=[
                  ["快速模式（默认）", "每个 PCIe 事务", "~1000-10000 事务/秒",
                   "驱动加载、功能验证、回归"],
                  ["精确模式", "每 N 个时钟周期（锁步）", "~10-100 事务/秒",
                   "时序调试、波形分析、性能测量"],
              ])

    add_heading(doc, "5.1 运行时切换", 2)
    add_para(doc, "通过 Bridge API 切换：")
    add_code(doc,
             "#include \"bridge_qemu.h\"\n\n"
             "// 切到精确模式\n"
             "bridge_request_mode_switch(ctx, COSIM_MODE_PRECISE);\n\n"
             "// 在精确模式下主动推进 N 个周期\n"
             "bridge_advance_clock(ctx, 1000);\n\n"
             "// 切回快速模式\n"
             "bridge_request_mode_switch(ctx, COSIM_MODE_FAST);")

    add_heading(doc, "5.2 事务追踪", 2)
    add_para(doc, "开启后，Bridge 会把 TLP / CPL / DMA / MSI 事件按顺序写入日志：")
    add_code(doc,
             "// CSV 格式\n"
             "bridge_enable_trace(ctx, \"/tmp/cosim_trace.csv\", TRACE_FMT_CSV);\n\n"
             "// 或 JSON 格式\n"
             "bridge_enable_trace(ctx, \"/tmp/cosim_trace.json\", TRACE_FMT_JSON);\n\n"
             "// 结束时关闭\n"
             "bridge_disable_trace(ctx);")
    add_para(doc, "CSV 头：timestamp,kind,type,tag,addr,len,data")
    add_para(doc, "JSON 为记录数组，每条包含 timestamp / kind / type / tag / addr / len / data。")

    doc.add_page_break()

    # ---------- 6. API 参考 ----------
    add_heading(doc, "6. API 参考", 1)

    add_heading(doc, "6.1 QEMU 侧 C API（bridge_qemu.h）", 2)
    add_table(doc,
              headers=["函数", "说明"],
              rows=[
                  ["bridge_init(shm_name, sock_path)",
                   "创建 SHM 并监听 socket，返回 bridge_ctx_t*"],
                  ["bridge_connect(ctx)",
                   "accept() VCS 连接，阻塞直到 VCS 接入"],
                  ["bridge_send_tlp(ctx, req)",
                   "排队 TLP 请求并通知 VCS"],
                  ["bridge_wait_completion(ctx, tag, cpl)",
                   "阻塞等待指定 tag 的 completion"],
                  ["bridge_send_tlp_and_wait(ctx, req, cpl)",
                   "上述两步的封装"],
                  ["bridge_complete_dma(ctx, tag, status)",
                   "对 VCS 发来的 DMA 请求应答完成"],
                  ["bridge_request_mode_switch(ctx, target_mode)",
                   "请求切换到 COSIM_MODE_FAST / COSIM_MODE_PRECISE"],
                  ["bridge_get_mode(ctx)",
                   "读取当前模式"],
                  ["bridge_advance_clock(ctx, cycles)",
                   "精确模式下推进 N 个周期并等待 VCS ACK"],
                  ["bridge_enable_trace(ctx, path, fmt)",
                   "开启事务追踪，fmt = CSV / JSON"],
                  ["bridge_disable_trace(ctx)",
                   "关闭追踪，flush 并关闭文件"],
                  ["bridge_destroy(ctx)",
                   "释放资源"],
              ])

    add_heading(doc, "6.2 VCS 侧 DPI-C 函数（bridge_vcs.sv）", 2)
    add_code(doc,
             "import \"DPI-C\" function int  vcs_bridge_init_dpi(\n"
             "    input string shm_name, input string sock_path);\n"
             "import \"DPI-C\" function int  vcs_bridge_poll_tlp_dpi(\n"
             "    output byte type_, output byte tag,\n"
             "    output shortint len, output longint addr,\n"
             "    output byte data[64]);\n"
             "import \"DPI-C\" function int  vcs_bridge_send_completion_dpi(\n"
             "    input byte tag, input shortint len, input byte data[64]);\n"
             "import \"DPI-C\" function int  vcs_bridge_trigger_dma_dpi(\n"
             "    input byte direction, input longint host_addr,\n"
             "    input int len, input int dma_offset);\n"
             "import \"DPI-C\" function int  vcs_bridge_raise_msi_dpi(input int vector);\n"
             "import \"DPI-C\" function int  vcs_bridge_clock_ack_dpi();\n"
             "import \"DPI-C\" function void vcs_bridge_close_dpi();")

    add_heading(doc, "6.3 QEMU 设备命令行参数", 2)
    add_code(doc,
             "-device cosim-pcie-rc,shm_name=/cosim0,sock_path=/tmp/cosim.sock")
    add_bullet(doc, "shm_name：POSIX 共享内存对象名（必填，以 / 开头）")
    add_bullet(doc, "sock_path：Unix Domain Socket 路径（必填）")

    doc.add_page_break()

    # ---------- 7. 开发指南 ----------
    add_heading(doc, "7. 开发指南", 1)

    add_heading(doc, "7.1 新增 PCIe 事务类型", 2)
    add_para(doc, "步骤：")
    add_bullet(doc, "在 bridge/common/cosim_types.h 的 tlp_type_t 枚举中添加新类型")
    add_bullet(doc, "若携带数据超过 64 字节，规划 DMA 区偏移（dma_offset 字段）")
    add_bullet(doc, "QEMU 侧：在 cosim_pcie_rc.c 的 cosim_mmio_read/write 或 DMA 回调中填充")
    add_bullet(doc, "VCS 侧：在 bridge_vcs.sv 的 poll/driver 中识别新类型")
    add_bullet(doc, "补 test_bridge_loopback 扩展测试")

    add_heading(doc, "7.2 新增 DMA 场景", 2)
    add_para(doc, "DMA 区（63 MB）使用 bump allocator（dma_manager.c）：")
    add_code(doc,
             "uint32_t off = dma_region_alloc(&shm, /*len=*/4096, /*align=*/64);\n"
             "if (off == DMA_ALLOC_FAIL) { /* 容量不足 */ }\n"
             "// 使用 (uint8_t*)shm.dma_buf + off 访问 4KB 区\n"
             "// Reset/fini 时 dma_region_reset(&shm)")
    add_note(doc, "Bump allocator 是顺序分配、整体重置模型。不支持单独释放。"
                  "如需复杂分配，建议在上层封装 free-list。")

    add_heading(doc, "7.3 运行单独一个测试", 2)
    add_code(doc,
             "cd build\n"
             "ctest -R test_dma_roundtrip -V")

    add_heading(doc, "7.4 扩展 E2E 测试", 2)
    add_para(doc, "当前 E2E 测试需要真实 VCS 环境。参考顺序：")
    add_bullet(doc, "准备 Guest 内核 + 驱动（lspci 能识别 VendorID 0x1234 DeviceID 0x0001）")
    add_bullet(doc, "启动 QEMU：./scripts/run_cosim.sh")
    add_bullet(doc, "启动 VCS：./simv +SHM_NAME=/cosim0 +SOCK_PATH=/tmp/cosim.sock")
    add_bullet(doc, "在 Guest 中：insmod driver.ko 触发 BAR 读写 / DMA / MSI")
    add_bullet(doc, "验证：/proc/interrupts 看到 MSI 计数增加、driver log 看到 DMA 完成")

    doc.add_page_break()

    # ---------- 8. 故障排查 ----------
    add_heading(doc, "8. 故障排查", 1)
    add_table(doc,
              headers=["现象", "可能原因", "排查步骤"],
              rows=[
                  ["bridge_init 失败，shm_open 报错",
                   "/dev/shm 空间不足或权限问题",
                   "df -h /dev/shm; 清理 /dev/shm/cosim* 旧对象"],
                  ["bridge_connect 卡住",
                   "VCS 未启动或 sock_path 不一致",
                   "确认两侧 sock_path 相同；ls -l /tmp/cosim.sock"],
                  ["MMIO 读写卡死",
                   "VCS 侧未处理 TLP / DPI-C 调用被阻塞",
                   "在 bridge_vcs.sv 的 poll 循环打印 tag；ring_buf_dequeue 是否返回 0"],
                  ["ctest 部分测试偶发超时",
                   "CI 机器负载过高 / 共享内存未清理",
                   "rm -f /dev/shm/cosim*; 单独重跑"],
                  ["精确模式下 advance_clock 返回 -1",
                   "当前不在 precise 模式",
                   "先 bridge_request_mode_switch(ctx, COSIM_MODE_PRECISE)"],
                  ["QEMU 启动报错 cosim: bridge_init failed",
                   "SHM 名被占用或权限",
                   "rm -f /dev/shm/cosim0; 重新启动"],
                  ["MSI 中断在 Guest 不可见",
                   "Guest 驱动未使能 MSI / QEMU 未加载 msi-x",
                   "lspci -vv 查看 MSI Enable；dmesg 查看 PCI 枚举日志"],
              ])

    add_heading(doc, "8.1 日志与调试", 2)
    add_bullet(doc, "QEMU 日志：-d guest_errors,unimp 可开启 cosim 模块日志")
    add_bullet(doc, "开启 trace：bridge_enable_trace() 写 /tmp/cosim_trace.csv，导入 Excel 分析")
    add_bullet(doc, "GDB 调试 Guest：run_cosim.sh 里 GDB=1 启动后在另一终端 gdb -ex 'target remote :1234'")
    add_bullet(doc, "SHM 布局 dump：调用 shm_layout.c 中的 cosim_shm_dump()（需自行添加）")

    doc.add_page_break()

    # ---------- 9. 发布与 CI ----------
    add_heading(doc, "9. 版本发布与 CI", 1)
    add_heading(doc, "9.1 已完成阶段", 2)
    add_table(doc,
              headers=["阶段", "交付内容", "状态"],
              rows=[
                  ["P1", "单节点 PCIe MMIO 通路、快速模式", "已完成，10 commits，test_bridge_loopback 通过"],
                  ["P2", "DMA + MSI + 精确模式 + Trace", "已完成，10 commits，5 集成测试通过"],
                  ["P3", "双节点 ETH 互打 + 链路模型 + launch_dual", "已完成，7 commits，4 集成测试通过"],
                  ["P4", "cosim_cli, trace_analyzer, GDB 文档, CI, smoke", "已完成，6 commits，CI Run 全绿"],
              ])

    add_heading(doc, "9.2 推送到远程仓库", 2)
    add_code(doc,
             "git remote -v\n"
             "# origin  https://github.com/Beihang-yuting/QEMU_VCS.git\n\n"
             "git push origin master")
    add_note(doc, "若使用 HTTPS，需要 Personal Access Token（带 repo scope）；"
                  "推荐配置 SSH 公钥避免每次输入。")

    add_heading(doc, "9.3 下一步计划", 2)
    add_bullet(doc, "补齐 README.md / LICENSE / .gitignore")
    add_bullet(doc, "P3：双节点以太网互打（bridge/eth + 链路模型）")
    add_bullet(doc, "P4：cosim_cli 调试控制台 + trace_analyzer 分析工具")
    add_bullet(doc, "CI：GitHub Actions 跑 make test；VCS 侧 smoke test 留为 manual trigger")

    doc.add_page_break()

    # ---------- 11. P3 双节点 ETH ----------
    add_heading(doc, "11. 双节点 ETH 互打（P3）", 1)

    add_para(doc,
             "P3 增加了第二条共享内存通路（ETH SHM），允许两个 CoSim 节点（各自一对 "
             "QEMU+VCS）通过一条软件以太网链路互通。链路模型支持丢包、突发丢包、固定延迟、"
             "速率限和流控窗口；时间同步采用松耦合模式（每节点事件驱动，事件戳通过 SHM "
             "barrier 单调推进）。")

    add_heading(doc, "11.1 双节点架构", 2)
    add_bullet(doc, "Node A：QEMU-A + VCS-A，通过 PCIe SHM 互联")
    add_bullet(doc, "Node B：QEMU-B + VCS-B，通过 PCIe SHM 互联")
    add_bullet(doc, "ETH SHM：A 与 B 节点之间的双向帧队列（A→B 和 B→A）")
    add_bullet(doc, "MAC stub（mac_stub.{c,h}）：本地软件 MAC，无需真实 RTL 即可端到端")
    add_bullet(doc, "DPI-C 接口（eth_mac_dpi.{c,h}）：真实 RTL MAC 上线后从此接入")

    add_heading(doc, "11.2 启动双节点（launch_dual.py）", 2)
    add_code(doc,
             "# 冒烟（不需要真实 QEMU/VCS）\n"
             "python3 scripts/launch_dual.py --launcher-cmd \"sleep 10\" --smoke\n\n"
             "# 真实双节点（需要 run_cosim.sh 已配置好 QEMU 镜像）\n"
             "python3 scripts/launch_dual.py \\\n"
             "    --shm-pcie-a /cosim-pcie-a --sock-a /tmp/cosim-a.sock \\\n"
             "    --shm-pcie-b /cosim-pcie-b --sock-b /tmp/cosim-b.sock \\\n"
             "    --shm-eth   /cosim-eth0\n\n"
             "# SSH 远程模式（节点 B 跑在另一台机器上）\n"
             "python3 scripts/launch_dual.py --mode ssh --node-b-host bob@dev.lan")

    add_heading(doc, "11.3 链路模型配置", 2)
    add_table(doc,
              headers=["字段", "类型", "示例", "说明"],
              rows=[
                  ["drop_rate_ppm",   "u32", "100000",  "每百万帧的进入丢包概率（10%）"],
                  ["burst_drop_len",  "u16", "5",       "一旦触发，连续丢的帧数"],
                  ["latency_ns",      "u64", "5000",    "每帧固定单向延迟（5us）"],
                  ["rate_mbps",       "u32", "1000",    "线速率，0 = 无限"],
                  ["fc_window",       "u32", "4",       "最大在飞帧数（流控），0 = 无限"],
              ])
    add_code(doc,
             "/* C 侧示例：模拟 1Gbps 链路 + 10% 丢包 + 4 帧 FC 窗口 */\n"
             "eth_port_t port = {0};\n"
             "port.link.drop_rate_ppm  = 100000;\n"
             "port.link.burst_drop_len = 1;\n"
             "port.link.latency_ns     = 5000;\n"
             "port.link.rate_mbps      = 1000;\n"
             "port.link.fc_window      = 4;\n"
             "eth_port_open(&port, \"/cosim-eth0\", ETH_ROLE_A, 1);")

    add_heading(doc, "11.4 ETH API（eth_port.h）", 2)
    add_table(doc,
              headers=["函数", "说明"],
              rows=[
                  ["eth_port_open(port, name, role, create)",
                   "打开 ETH 端口，绑定 SHM 与角色（A/B）"],
                  ["eth_port_send(port, frame, now_ns)",
                   "发帧；返回 0 / -1 ring full / -2 FC 阻塞 / -3 链路丢包"],
                  ["eth_port_recv(port, out, timeout_ns)",
                   "收帧；timeout=0 非阻塞，>0 阻塞超时"],
                  ["eth_port_tx_complete(port)",
                   "对端已消费 → 减 outstanding（FC 用）"],
                  ["eth_port_close(port)",
                   "关闭端口，owned_shm=1 时自动 unlink SHM"],
                  ["eth_shm_advance_time(shm, role, ns)",
                   "松耦合：发布本节点 sim 时间"],
                  ["eth_shm_peer_time(shm, self_role)",
                   "松耦合：读取对方节点最近时间"],
              ])

    add_heading(doc, "11.5 VCS RTL MAC 接入路径", 2)
    add_para(doc, "当真实 MAC RTL 就绪时，按以下步骤替换 MAC stub：")
    add_bullet(doc, "在 SystemVerilog 顶层 import \"DPI-C\" 声明 vcs_eth_mac_*_dpi 函数")
    add_bullet(doc, "RTL MAC TX：MAC 输出帧后调 vcs_eth_mac_send_frame_dpi(data, len)")
    add_bullet(doc, "RTL MAC RX 轮询：每周期或空闲调 vcs_eth_mac_poll_frame_dpi(buf, max_len)")
    add_bullet(doc, "复位 / 关闭：vcs_eth_mac_close_dpi()")
    add_bullet(doc, "链路参数运行时调整：vcs_eth_mac_configure_link_dpi(...)")

    doc.add_page_break()

    # ---------- 12. 调试工具（P4）----------
    add_heading(doc, "12. 调试工具（P4）", 1)

    add_heading(doc, "12.1 cosim_cli 交互式控制台", 2)
    add_code(doc,
             "# 启动（先 make bridge）\n"
             "python3 scripts/cosim_cli.py --shm /cosim0 --sock /tmp/cosim.sock\n\n"
             "# REPL 内常用命令\n"
             "cosim> read 0x100\n"
             "cosim> write 0x200 0xDEADBEEF 4\n"
             "cosim> mode precise\n"
             "cosim> advance 1000\n"
             "cosim> trace on /tmp/trace.csv csv\n"
             "cosim> status\n"
             "cosim> quit")

    add_heading(doc, "12.2 trace_analyzer 事务分析", 2)
    add_code(doc,
             "# 解析 CSV 或 JSON trace 文件\n"
             "python3 scripts/trace_analyzer.py /tmp/trace.csv\n\n"
             "# 输出包含：\n"
             "#   - 事件分布（按 kind / type）\n"
             "#   - tag 匹配统计 + 孤立 cpl 检测\n"
             "#   - 延迟统计（min/mean/median/p95/p99）\n"
             "#   - ASCII 直方图\n"
             "#   - DMA 总字节、MSI 向量分布")

    add_heading(doc, "12.3 GDB 调试", 2)
    add_para(doc, "完整指南：docs/GDB-Debugging-Guide.md")
    add_code(doc,
             "# Guest 内核：\n"
             "GDB=1 ./scripts/run_cosim.sh        # QEMU 在 :1234 等待\n"
             "gdb vmlinux\n"
             "(gdb) target remote :1234\n\n"
             "# Bridge 库：\n"
             "gdb build/tests/integration/test_dma_roundtrip\n"
             "(gdb) break bridge_complete_dma")

    add_heading(doc, "12.4 CI / actionlint", 2)
    add_bullet(doc, ".github/workflows/ci.yml：每次 push/PR 自动 cmake build + ctest")
    add_bullet(doc, "Lint job：shellcheck + python compile 检查")
    add_bullet(doc, "本地校验：actionlint -no-color .github/workflows/ci.yml")

    doc.add_page_break()

    # ---------- 10. 附录 ----------
    add_heading(doc, "10. 附录", 1)

    add_heading(doc, "10.1 共享内存布局", 2)
    add_table(doc,
              headers=["偏移", "大小", "内容"],
              rows=[
                  ["0x0000_0000", "4 KB", "控制区：magic, version, mode, ready, sim_time, IRQ 状态"],
                  ["0x0000_1000", "256 KB", "请求队列（Host→Device）：TLP 环形缓冲"],
                  ["0x0004_1000", "256 KB", "响应队列（Device→Host）：Completion 环形缓冲"],
                  ["0x0008_1000", "64 KB", "DMA 请求队列 + DMA 完成队列 + MSI 事件队列"],
                  ["0x0009_1000", "~63 MB", "DMA 数据缓冲区（bump allocator）"],
              ])

    add_heading(doc, "10.2 PCI ID 分配", 2)
    add_bullet(doc, "Vendor ID: 0x1234（平台自用保留）")
    add_bullet(doc, "Device ID: 0x0001（CoSim PCIe RC）")
    add_bullet(doc, "Class:     0x0200（Ethernet Controller，便于现有驱动加载）")

    add_heading(doc, "10.3 消息类型（cosim_types.h）", 2)
    add_code(doc,
             "TLP_MWR, TLP_MRD, TLP_CFGWR, TLP_CFGRD, TLP_CPL\n"
             "SYNC_MSG_TLP_READY, SYNC_MSG_CPL_READY,\n"
             "SYNC_MSG_DMA_REQ, SYNC_MSG_DMA_CPL,\n"
             "SYNC_MSG_MSI,     SYNC_MSG_MODE_SWITCH,\n"
             "SYNC_MSG_CLOCK_STEP, SYNC_MSG_CLOCK_ACK")

    add_heading(doc, "10.4 相关文档", 2)
    add_bullet(doc, "设计文档：docs/superpowers/specs/2026-04-16-qemu-vcs-cosim-platform-design.md")
    add_bullet(doc, "P1 实施计划：docs/superpowers/plans/2026-04-16-cosim-p1-pcie-path.md")
    add_bullet(doc, "P2 实施计划：docs/superpowers/plans/2026-04-16-cosim-p2-dma-interrupt-precise.md")
    add_bullet(doc, "PPT（架构总览）：QEMU-VCS-CoSim-Platform-Design.pptx（项目根目录）")

    add_heading(doc, "10.5 术语表", 2)
    add_table(doc,
              headers=["缩写 / 术语", "含义"],
              rows=[
                  ["TLP", "Transaction Layer Packet，PCIe 事务层包"],
                  ["MMIO", "Memory-Mapped I/O，通过内存读写访问设备寄存器"],
                  ["DMA", "Direct Memory Access，设备直接读写主机内存"],
                  ["MSI", "Message Signaled Interrupts，消息式中断"],
                  ["SHM", "POSIX Shared Memory，进程间共享内存"],
                  ["DPI-C", "Direct Programming Interface for C，VCS 调 C 的桥"],
                  ["BAR", "Base Address Register，PCIe 配置空间地址寄存器"],
                  ["EP", "Endpoint，PCIe 端点设备（如 DPU）"],
                  ["RC", "Root Complex，PCIe 根复合体（Host 侧）"],
              ])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Generated: {OUTPUT}")
    print(f"Size:      {os.path.getsize(OUTPUT)} bytes")


if __name__ == "__main__":
    build()
